"""document_tool.py — Preview/inspection utilities for a filled .docx.

Used by nodes/human_review.py to let the student see what was actually
filled in a readable chat-friendly format, WITHOUT having to download
and open the .docx first. The actual file is only offered for download
after the student has reviewed and confirmed it looks right.
"""

from __future__ import annotations

import io


def extract_preview_text(docx_bytes: bytes, max_chars: int = 3000) -> str:
    """Extract readable text (paragraphs + table cells) from a filled
    .docx for an in-chat preview.

    This is intentionally a light, best-effort text dump — not a
    full-fidelity rendering (no formatting, no layout) — good enough
    for "does this look right before I download it", not meant to
    replace opening the real file.
    """
    from docx import Document

    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001
        return f"(Không đọc được bản xem trước: {exc})"

    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            # Collapse repeated identical cells from merged-cell
            # duplication (see form_filler.py's fill notes) so the
            # preview doesn't show the same value 3-4 times per row.
            deduped = list(dict.fromkeys(t for t in cell_texts if t))
            if deduped:
                lines.append(" | ".join(deduped))

    full_text = "\n".join(lines)

    if len(full_text) > max_chars:
        return full_text[:max_chars].rstrip() + "\n\n… (đã cắt bớt, xem đầy đủ trong file tải về)"

    return full_text


def document_byte_size_kb(docx_bytes: bytes) -> float:
    """Small helper for showing file size in the review UI."""
    return round(len(docx_bytes) / 1024, 1)