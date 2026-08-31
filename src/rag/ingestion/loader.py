"""loader.py — Document loading (PDF, DOCX).

Extracts text and structure from PDF and DOCX files.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from pypdf import PdfReader


WHITESPACE_RE = re.compile(r"\s+")


DOCUMENT_RULES = (
    ("POL-CAID", "policy"),
    ("Internship-Management-Policy", "policy"),
    ("Form-1", "form"),
    ("Form-2", "agreement"),
    ("Form-3", "form"),
    ("Form-4", "form"),
    ("Talent-Handbook", "talent_handbook"),
    ("Capstone-Booklet", "capstone_booklet"),
)


SEMANTIC_DOCUMENT_TYPES = frozenset({
    "policy",
    "form",
    "agreement",
    "talent_handbook",
    "capstone_booklet",
    "knowledge",
})


@dataclass
class ExtractedElement:
    document_name: str
    document_type: str
    file_type: str
    text: str
    page: int | None = None
    element_type: str | None = None
    element_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "document_name": self.document_name,
            "document_type": self.document_type,
            "file_type": self.file_type,
            "text": self.text,
            "page": self.page,
            "element_type": self.element_type,
            "element_index": self.element_index,
            "table_index": self.table_index,
            "row_index": self.row_index,
        }


@dataclass
class ExtractionResult:
    document_name: str
    document_type: str
    file_type: str
    elements: list[ExtractedElement] = field(default_factory=list)
    pages: int = 0
    paragraphs: int = 0
    tables: int = 0
    characters_extracted: int = 0
    empty_pages: list[int] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    @property
    def status(self) -> str:
        if self.requires_ocr:
            return "requires_ocr"
        if self.elements and not self.errors:
            return "success"
        if self.elements:
            return "partial"
        return "failed"

    @property
    def requires_ocr(self) -> bool:
        return (
            self.file_type == "pdf"
            and self.pages > 0
            and not self.elements
            and any("require OCR" in error for error in self.errors)
        )

    def all_text(self) -> str:
        return "\n".join(element.text for element in self.elements if element.text)

    def report(self) -> dict[str, Any]:
        return {
            "document_name": self.document_name,
            "document_type": self.document_type,
            "file_type": self.file_type,
            "pages": self.pages,
            "paragraphs": self.paragraphs,
            "tables": self.tables,
            "characters_extracted": self.characters_extracted,
            "empty_pages": self.empty_pages,
            "errors": self.errors,
            "status": self.status,
        }


def normalize_text(text: str) -> str:
    return WHITESPACE_RE.sub(" ", text.replace("\x00", " ")).strip()


def classify_document(path: Path) -> str:
    """Legacy filename-based semantic classification."""
    name = path.name
    for pattern, document_type in DOCUMENT_RULES:
        if pattern.lower() in name.lower():
            return document_type
    return "unknown"


def resolve_document_type(
    path: Path,
    document_type_override: str | None = None,
) -> str:
    """Resolve the semantic RAG document type.

    Admin-managed documents should provide ``document_type_override``.
    Filename classification remains only as a backwards-compatible fallback
    for the original built-in RAG corpus.
    """
    if document_type_override is None:
        return classify_document(path)

    normalized = str(document_type_override).strip().lower()

    if normalized not in SEMANTIC_DOCUMENT_TYPES:
        allowed = ", ".join(sorted(SEMANTIC_DOCUMENT_TYPES))
        raise ValueError(
            f"Unsupported semantic document type: "
            f"{document_type_override!r}. Allowed: {allowed}"
        )

    return normalized


def load_document(
    path: Path,
    document_type_override: str | None = None,
) -> ExtractionResult:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(
            path,
            document_type_override=document_type_override,
        )

    if suffix == ".docx":
        return load_docx(
            path,
            document_type_override=document_type_override,
        )

    result = ExtractionResult(
        document_name=path.name,
        document_type=resolve_document_type(
            path,
            document_type_override,
        ),
        file_type=suffix.lstrip(".") or "unknown",
    )
    result.errors.append(f"Unsupported file type: {suffix}")
    return result


def load_pdf(
    path: Path,
    document_type_override: str | None = None,
) -> ExtractionResult:
    result = ExtractionResult(
        document_name=path.name,
        document_type=resolve_document_type(
            path,
            document_type_override,
        ),
        file_type="pdf",
    )
    _load_pdf_with_pypdf(path, result)
    if result.elements:
        return result

    # Some designed PDFs expose no text through pypdf even though text is
    # selectable. PyMuPDF gives us a second extraction path before we classify
    # pages as image-only/empty.
    pypdf_errors = list(result.errors)
    result.elements.clear()
    result.characters_extracted = 0
    result.empty_pages.clear()
    result.errors.clear()
    _load_pdf_with_pymupdf(path, result)
    result.errors = pypdf_errors + result.errors
    if result.pages and not result.elements:
        result.errors.append(
            "No extractable text found with pypdf or PyMuPDF; the PDF may be image-only and require OCR."
        )
    return result


def _load_pdf_with_pypdf(path: Path, result: ExtractionResult) -> None:
    try:
        reader = PdfReader(str(path))
        result.pages = len(reader.pages)
        for page_index, page in enumerate(reader.pages, start=1):
            try:
                text = normalize_text(page.extract_text() or "")
            except Exception as exc:  # pypdf can fail on individual pages.
                result.errors.append(f"page {page_index}: {exc}")
                text = ""

            if not text:
                result.empty_pages.append(page_index)
                continue

            result.elements.append(
                ExtractedElement(
                    document_name=path.name,
                    document_type=result.document_type,
                    file_type="pdf",
                    page=page_index,
                    element_type="page",
                    element_index=page_index,
                    text=text,
                )
            )
            result.characters_extracted += len(text)
    except Exception as exc:
        result.errors.append(str(exc))


def _load_pdf_with_pymupdf(path: Path, result: ExtractionResult) -> None:
    try:
        with fitz.open(str(path)) as document:
            result.pages = document.page_count
            for page_index, page in enumerate(document, start=1):
                try:
                    text = normalize_text(page.get_text("text") or "")
                except Exception as exc:
                    result.errors.append(f"page {page_index}: {exc}")
                    text = ""

                if not text:
                    result.empty_pages.append(page_index)
                    continue

                result.elements.append(
                    ExtractedElement(
                        document_name=path.name,
                        document_type=result.document_type,
                        file_type="pdf",
                        page=page_index,
                        element_type="page",
                        element_index=page_index,
                        text=text,
                    )
                )
                result.characters_extracted += len(text)
    except Exception as exc:
        result.errors.append(f"pymupdf: {exc}")


def load_docx(
    path: Path,
    document_type_override: str | None = None,
) -> ExtractionResult:
    result = ExtractionResult(
        document_name=path.name,
        document_type=resolve_document_type(
            path,
            document_type_override,
        ),
        file_type="docx",
    )
    try:
        document = Document(str(path))
    except Exception as exc:
        result.errors.append(str(exc))
        return result

    element_index = 0
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        element_index += 1
        result.paragraphs += 1
        style_name = getattr(paragraph.style, "name", "") or ""
        element_type = "heading" if style_name.lower().startswith("heading") else "paragraph"
        result.elements.append(
            ExtractedElement(
                document_name=path.name,
                document_type=result.document_type,
                file_type="docx",
                element_type=element_type,
                element_index=element_index,
                text=text,
            )
        )
        result.characters_extracted += len(text)

    for table_index, table in enumerate(document.tables, start=1):
        result.tables += 1
        for row_index, row in enumerate(table.rows, start=1):
            cells = [normalize_text(cell.text) for cell in row.cells]
            # Repeated merged-cell text is common in DOCX tables; preserve order
            # while removing adjacent duplicates to keep reports readable.
            deduped_cells: list[str] = []
            for cell in cells:
                if cell and (not deduped_cells or deduped_cells[-1] != cell):
                    deduped_cells.append(cell)
            text = normalize_text(" | ".join(deduped_cells))
            if not text:
                continue
            element_index += 1
            result.elements.append(
                ExtractedElement(
                    document_name=path.name,
                    document_type=result.document_type,
                    file_type="docx",
                    element_type="table_row",
                    element_index=element_index,
                    table_index=table_index,
                    row_index=row_index,
                    text=text,
                )
            )
            result.characters_extracted += len(text)

    return result
