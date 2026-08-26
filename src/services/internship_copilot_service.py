from __future__ import annotations

import hashlib
import io
import json
import logging
import math
import re
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from typing import Any, Callable, Literal
from zoneinfo import ZoneInfo

from pydantic import BaseModel, Field
from sqlalchemy import text
from sqlalchemy.orm import Session

from src.config import get_settings
from src.observability.instrumentation import langfuse_callbacks
from src.rag.query_pipeline import (
    ALL_ROUTED_DOCUMENT_TYPES,
    CAREER_DOCUMENT_TYPES,
    INTERNSHIP_DOCUMENT_TYPES,
    RouteDecision,
    _get_chat_llm,
)
from src.rag.schemas import QueryResult
from src.services.notification_service import (
    get_pending_reminders,
    schedule_reminder,
)

logger = logging.getLogger(__name__)

RagLookup = Callable[[str, RouteDecision], QueryResult]

COPILOT_DB_ACTIONS = {
    "eligibility_checker",
    "internship_checklist",
    "deadline_timeline",
    "internship_matching",
    "cv_improvement",
    "jd_analyzer",
    "interview_preparation",
    "internship_progress",
    "weekly_reflection",
    "evaluation_preparation",
    "skill_gap_analysis",
    "career_recommendation",
    "smart_notifications",
    "personalized_dashboard",
    "human_escalation",
    "grievance_assistant",
}

WRITE_ACTIONS = {
    "internship_progress",
    "weekly_reflection",
    "smart_notifications",
    "human_escalation",
    "grievance_assistant",
}


@dataclass
class StudentSnapshot:
    profile: dict[str, Any] | None
    internship: dict[str, Any] | None
    application: dict[str, Any] | None


class EscalationDraft(BaseModel):
    # Semantic completeness gate. This is deliberately separate from the
    # requested operation/language/recipient so commands such as
    # "tạo escalation bằng tiếng Việt" are NOT converted into fake incidents.
    has_incident_details: bool = False
    clarification_question: str | None = None
    escalation_type: Literal[
        "SAFETY",
        "SUPERVISION",
        "WORKLOAD",
        "HARASSMENT",
        "ROLE_MISMATCH",
        "WITHDRAWAL",
        "OTHER",
    ] = "OTHER"
    severity: Literal["LOW", "MEDIUM", "HIGH", "CRITICAL"] = "MEDIUM"
    target: Literal["FACULTY_MENTOR", "CAID_QUEUE"] = "FACULTY_MENTOR"
    subject: str | None = Field(default=None, max_length=255)
    description: str | None = Field(default=None, max_length=4000)


class ProgressWriteDraft(BaseModel):
    has_work_details: bool = False
    clarification_question: str | None = None
    work_summary: str | None = Field(default=None, max_length=4000)
    hours: float | None = Field(default=None, ge=0, le=168)
    week: int | None = Field(default=None, ge=1, le=60)


class ReflectionWriteDraft(BaseModel):
    has_target_week: bool = False
    clarification_question: str | None = None
    week: int | None = Field(default=None, ge=1, le=60)


class ReminderWriteDraft(BaseModel):
    kind: Literal["REMINDER", "PREFERENCE"] = "REMINDER"
    has_required_details: bool = False
    clarification_question: str | None = None
    reminder_content: str | None = Field(default=None, max_length=1000)
    time_expression: str | None = Field(default=None, max_length=200)
    scheduled_at_iso: str | None = Field(default=None, max_length=100)
    days_before: int | None = Field(default=None, ge=0, le=365)
    deadline_reference: str | None = Field(default=None, max_length=255)
    preference_key: Literal[
        "report_deadline",
        "lecturer_feedback",
        "internship_status",
        "email_notifications",
    ] | None = None
    preference_enabled: bool | None = None


class ExtractedSkills(BaseModel):
    skills: list[str] = Field(default_factory=list)


# ---------------------------------------------------------------------------
# Public dispatch
# ---------------------------------------------------------------------------

def handle_internship_copilot_action(
    db: Session,
    current_user: dict,
    message: str,
    route: RouteDecision,
    session_id: str | None = None,
    rag_lookup: RagLookup | None = None,
    conversation_context: str = "",
) -> QueryResult | None:
    """Handle authenticated Internship Copilot actions that need DB/tools.

    The shared semantic router remains authoritative. This layer never opens the
    personal database based on keyword matching: DB-backed analysis only runs for
    an authenticated STUDENT and a route that either explicitly selected personal
    scope or selected an explicit write action.
    """
    action = getattr(route, "assistant_action", "none") or "none"
    action_mode = getattr(route, "action_mode", "inform") or "inform"
    data_source = getattr(route, "data_source", "none") or "none"

    # Defense in depth: an action name never authorizes DB/write access by itself.
    if action_mode == "inform" and data_source != "personal_db":
        return None
    if action_mode == "execute" and data_source != "write_action":
        return None

    if action not in COPILOT_DB_ACTIONS:
        return None

    if str(current_user.get("role") or "").upper() != "STUDENT":
        return None

    # Read access is fail-closed: only personal scope can use stored student data.
    # Explicit write actions are also allowed because their targets are necessarily
    # the authenticated student's own internship records.
    if route.scope != "personal" and not (
        action_mode == "execute" and action in WRITE_ACTIONS
    ):
        return None

    if route.needs_clarification:
        return None

    student_id = int(current_user["id"])
    language = (
        getattr(route, "response_language", None)
        if getattr(route, "response_language", None) in {"vi", "en"}
        else route.language
        if route.language in {"vi", "en"}
        else "vi"
    )

    requested_sections = set(getattr(route, "personal_sections", []) or [])
    requested_profile_fields = set(getattr(route, "personal_profile_fields", []) or [])

    # Access only the minimum DB sections needed by the selected action. We do
    # not load the whole student record merely because the user is authenticated.
    need_profile = action in {
        "eligibility_checker", "internship_matching", "skill_gap_analysis",
        "career_recommendation",
    } and (
        route.scope != "personal"
        or ("profile" in requested_sections and bool(requested_profile_fields))
    )
    need_internship = action in {
        "eligibility_checker", "internship_checklist", "deadline_timeline",
        "jd_analyzer", "interview_preparation", "internship_progress",
        "weekly_reflection", "evaluation_preparation", "skill_gap_analysis",
        "personalized_dashboard", "human_escalation",
        "grievance_assistant",
    } and (
        route.scope != "personal"
        or "internship" in requested_sections
        or action_mode == "execute"
    )
    need_application = action in {
        "eligibility_checker", "internship_checklist", "cv_improvement",
        "jd_analyzer", "interview_preparation", "skill_gap_analysis",
        "personalized_dashboard",
    } and (
        route.scope != "personal"
        or "applications" in requested_sections
    )

    snapshot = StudentSnapshot(
        profile=(
            _get_student_profile(
                db, student_id,
                requested_fields=(requested_profile_fields if route.scope == "personal" else {"major", "gpa", "skills"}),
            )
            if need_profile else None
        ),
        internship=_get_current_internship(db, student_id) if need_internship else None,
        application=_get_current_application(db, student_id) if need_application else None,
    )

    if action == "personalized_dashboard":
        return _dashboard_result(
            db, student_id, message, language, snapshot, allowed_sections=requested_sections
        )

    if action == "eligibility_checker":
        return _eligibility_result(
            db, student_id, message, language, snapshot, rag_lookup
        )

    if action == "internship_checklist":
        return _checklist_result(
            db, student_id, message, language, snapshot, rag_lookup,
            allowed_sections=requested_sections,
        )

    if action == "deadline_timeline":
        return _deadline_result(db, student_id, message, language, snapshot)

    if action == "internship_matching":
        return _matching_result(
            db, student_id, message, language, snapshot,
            use_stored_cv=(route.scope != "personal" or "documents" in requested_sections),
        )

    if action == "cv_improvement":
        if route.scope == "personal" and "documents" not in requested_sections:
            return None
        return _cv_improvement_result(
            db, student_id, message, language, snapshot,
            allow_application_context=("applications" in requested_sections),
        )

    if action == "jd_analyzer":
        if route.scope == "personal" and not ({"documents", "applications"} & requested_sections):
            return None
        return _jd_result(
            db, student_id, message, language, snapshot,
            allow_documents=(route.scope != "personal" or "documents" in requested_sections),
            allow_application_context=(route.scope != "personal" or "applications" in requested_sections),
        )

    if action == "skill_gap_analysis":
        if route.scope == "personal" and not ({"profile", "documents", "applications"} <= requested_sections):
            return None
        return _skill_gap_result(
            db, student_id, message, language, snapshot,
            allow_documents=(route.scope != "personal" or "documents" in requested_sections),
            allow_application_context=(route.scope != "personal" or "applications" in requested_sections),
        )

    if action == "interview_preparation":
        if route.scope == "personal" and not ({"documents", "applications"} <= requested_sections):
            return None
        return _interview_result(
            db, student_id, message, language, snapshot,
            allow_documents=(route.scope != "personal" or "documents" in requested_sections),
            allow_application_context=(route.scope != "personal" or "applications" in requested_sections),
        )

    if action == "career_recommendation":
        return _career_result(
            db, student_id, message, language, snapshot, rag_lookup
        )

    if action == "internship_progress":
        if action_mode == "execute":
            return _save_progress_result(db, student_id, message, language, snapshot, route=route, conversation_context=conversation_context)
        if route.scope == "personal" and "progress" not in requested_sections:
            return None
        return _progress_summary_result(db, student_id, message, language, snapshot)

    if action == "weekly_reflection":
        if (
            route.scope == "personal"
            and action_mode != "execute"
            and not ({"progress", "reports"} <= requested_sections)
        ):
            return None
        return _weekly_reflection_result(
            db,
            student_id,
            message,
            language,
            snapshot,
            save_draft=(action_mode == "execute"),
            conversation_context=(conversation_context if action_mode == "execute" else ""),
            route=(route if action_mode == "execute" else None),
        )

    if action == "evaluation_preparation":
        if route.scope == "personal" and "evaluations" not in requested_sections:
            return None
        return _evaluation_result(db, student_id, message, language, snapshot, rag_lookup)

    if action == "smart_notifications":
        if action_mode == "execute":
            return _schedule_reminder_result(db, student_id, message, language, snapshot, route=route, conversation_context=conversation_context)
        return _reminder_overview_result(
            db, student_id, message, language, snapshot,
            include_deadlines=(route.scope != "personal" or "deadlines" in requested_sections),
        )

    if action in {"human_escalation", "grievance_assistant"}:
        if action_mode == "execute":
            return _create_escalation_result(
                db, student_id, message, language, snapshot,
                conversation_context=conversation_context,
                route=route,
            )
        if route.scope == "personal" and "escalations" in set(getattr(route, "personal_sections", []) or []):
            return _escalation_overview_result(db, student_id, message, language)
        # Guidance itself should remain grounded in the official grievance/policy RAG path.
        return None

    return None



# ---------------------------------------------------------------------------
# Semantic write contracts — NO extra LLM call
# ---------------------------------------------------------------------------

def _route_action_payload(route: RouteDecision) -> dict[str, Any]:
    """Return the payload extracted by the existing semantic-router call."""
    payload = getattr(route, "action_payload", None)
    if payload is None:
        return {}
    if hasattr(payload, "model_dump"):
        try:
            return dict(payload.model_dump(mode="python"))
        except Exception:
            return dict(payload.model_dump())
    if isinstance(payload, dict):
        return dict(payload)
    return {}


def _route_missing_fields(route: RouteDecision) -> list[str]:
    return [str(x) for x in (getattr(route, "missing_action_fields", None) or []) if str(x).strip()]


def _draft_progress_write(
    route: RouteDecision,
    language: str,
) -> ProgressWriteDraft:
    question = (
        "Bạn muốn ghi nhận công việc/progress cụ thể nào? Hãy mô tả việc đã làm; số giờ là tùy chọn."
        if language == "vi" else
        "What specific work/progress would you like to record? Describe the work completed; hours are optional."
    )
    payload = _route_action_payload(route)
    summary = str(payload.get("progress_work_summary") or "").strip() or None
    hours = payload.get("progress_hours")
    week = payload.get("progress_week")
    complete = bool(summary) and "progress_work_summary" not in _route_missing_fields(route)
    return ProgressWriteDraft(
        has_work_details=complete,
        clarification_question=(None if complete else question),
        work_summary=summary if complete else None,
        hours=(float(hours) if hours is not None else None),
        week=(int(week) if week is not None else None),
    )


def _draft_reflection_write(
    route: RouteDecision,
    language: str,
) -> ReflectionWriteDraft:
    question = (
        "Bạn muốn lưu reflection của tuần nào?"
        if language == "vi" else
        "Which week would you like to save the reflection for?"
    )
    payload = _route_action_payload(route)
    week = payload.get("reflection_week")
    complete = week is not None and "reflection_week" not in _route_missing_fields(route)
    return ReflectionWriteDraft(
        has_target_week=complete,
        clarification_question=(None if complete else question),
        week=(int(week) if complete else None),
    )


def _draft_reminder_write(
    route: RouteDecision,
    language: str,
) -> ReminderWriteDraft:
    question = (
        "Bạn muốn mình nhắc việc gì và vào thời điểm nào?"
        if language == "vi" else
        "What would you like me to remind you about, and when?"
    )
    payload = _route_action_payload(route)
    kind = str(payload.get("reminder_kind") or "REMINDER").upper()
    if kind not in {"REMINDER", "PREFERENCE"}:
        kind = "REMINDER"

    if kind == "PREFERENCE":
        key = payload.get("notification_preference_key")
        enabled = payload.get("notification_preference_enabled")
        complete = key is not None and enabled is not None and not _route_missing_fields(route)
        return ReminderWriteDraft(
            kind="PREFERENCE",
            has_required_details=complete,
            clarification_question=(None if complete else (
                "Bạn muốn bật/tắt loại thông báo nào?"
                if language == "vi" else
                "Which notification category would you like to enable or disable?"
            )),
            preference_key=key,
            preference_enabled=(bool(enabled) if enabled is not None else None),
        )

    content = str(payload.get("reminder_content") or "").strip() or None
    time_expression = str(payload.get("reminder_time_expression") or "").strip() or None
    scheduled_at_iso = str(payload.get("reminder_scheduled_at") or "").strip() or None
    days_before = payload.get("reminder_days_before")
    deadline_reference = str(payload.get("reminder_deadline_reference") or "").strip() or None
    complete = bool(content) and (
        bool(scheduled_at_iso)
        or (days_before is not None and bool(deadline_reference))
    ) and not _route_missing_fields(route)
    return ReminderWriteDraft(
        kind="REMINDER",
        has_required_details=complete,
        clarification_question=(None if complete else question),
        reminder_content=(content if complete else content),
        time_expression=time_expression,
        scheduled_at_iso=scheduled_at_iso,
        days_before=(int(days_before) if days_before is not None else None),
        deadline_reference=deadline_reference,
    )


def _draft_escalation(
    route: RouteDecision,
    language: str,
) -> EscalationDraft:
    clarification = (
        "Bạn muốn escalation phản ánh vấn đề/sự việc gì trong kỳ thực tập?"
        if language == "vi"
        else "What internship issue or incident would you like the escalation to report?"
    )
    payload = _route_action_payload(route)
    incident = str(payload.get("escalation_incident_description") or "").strip() or None
    missing = set(_route_missing_fields(route))
    complete = bool(incident) and "escalation_incident_description" not in missing

    if not complete:
        return EscalationDraft(
            has_incident_details=False,
            clarification_question=clarification,
            target="FACULTY_MENTOR",
            subject=None,
            description=None,
        )

    subject = str(payload.get("escalation_subject") or "").strip()
    if not subject:
        # Deterministic, fact-preserving fallback: use a short prefix of the user's
        # incident itself rather than inventing a new event/claim.
        subject = incident[:120].strip()
        if len(incident) > 120:
            subject = subject.rstrip(" ,.;:-") + "…"

    escalation_type = str(payload.get("escalation_type") or "OTHER").upper()
    if escalation_type not in {
        "SAFETY", "SUPERVISION", "WORKLOAD", "HARASSMENT",
        "ROLE_MISMATCH", "WITHDRAWAL", "OTHER",
    }:
        escalation_type = "OTHER"

    severity = str(payload.get("escalation_severity") or "MEDIUM").upper()
    if severity not in {"LOW", "MEDIUM", "HIGH", "CRITICAL"}:
        severity = "MEDIUM"

    target = str(payload.get("escalation_target") or "FACULTY_MENTOR").upper()
    if target not in {"FACULTY_MENTOR", "CAID_QUEUE"}:
        target = "FACULTY_MENTOR"

    return EscalationDraft(
        has_incident_details=True,
        clarification_question=None,
        escalation_type=escalation_type,
        severity=severity,
        target=target,
        subject=subject,
        description=incident,
    )


def _clarification_result(
    message: str,
    language: str,
    action: str,
    question: str,
) -> QueryResult:
    return _result(message, question.strip(), language, action).model_copy(
        update={
            "guardrail_reason": "copilot_clarification_required",
            "groundedness_status": "skip",
            "groundedness_reason": "copilot_clarification_required",
        }
    )


# ---------------------------------------------------------------------------
# Confirmation-gated write previews (NO persistent side effects)
# ---------------------------------------------------------------------------

def preview_internship_copilot_action(
    db: Session,
    current_user: dict,
    message: str,
    route: RouteDecision,
    session_id: str | None = None,
    conversation_context: str = "",
) -> QueryResult | None:
    """Describe a supported write before execution.

    This function may read the authenticated student's current internship/deadline
    context so the preview is concrete, but it never INSERTs/UPDATEs/DELETEs or
    sends a message/notification. The API layer must require a later confirmation
    turn before calling the normal execute path.
    """
    action = getattr(route, "assistant_action", "none") or "none"
    if getattr(route, "data_source", "none") != "write_action":
        return None
    if action not in WRITE_ACTIONS:
        return None
    if str(current_user.get("role") or "").upper() != "STUDENT":
        return None

    # Do NOT drop a semantically selected write_action merely because the router
    # marked it incomplete. The preview builders below already validate the
    # structured action_payload and return the minimum clarification when needed.
    # Returning None here previously caused write requests to fall through into
    # the normal RAG/general-answer pipeline.
    student_id = int(current_user["id"])
    language = (
        getattr(route, "response_language", None)
        if getattr(route, "response_language", None) in {"vi", "en"}
        else route.language
        if route.language in {"vi", "en"}
        else "vi"
    )
    internship = _get_current_internship(db, student_id)
    snapshot = StudentSnapshot(profile=None, internship=internship, application=None)

    if action == "internship_progress":
        return _preview_progress_result(message, language, snapshot, route=route, conversation_context=conversation_context)
    if action == "weekly_reflection":
        return _preview_weekly_reflection_result(db, message, language, snapshot, route=route, conversation_context=conversation_context)
    if action == "smart_notifications":
        return _preview_reminder_result(db, student_id, message, language, snapshot, route=route, conversation_context=conversation_context)
    if action in {"human_escalation", "grievance_assistant"}:
        return _preview_escalation_result(
            message, language, snapshot, route=route, conversation_context=conversation_context
        )
    return None


def make_copilot_confirmation_result(
    message: str,
    language: str,
    *,
    action: str,
    state: Literal["cancelled", "no_pending", "failed"],
) -> QueryResult:
    if language not in {"vi", "en"}:
        language = "vi"
    if state == "cancelled":
        answer = (
            "Đã hủy thao tác đang chờ xác nhận. **Không có thay đổi nào được thực hiện.**"
            if language == "vi" else
            "The pending action was cancelled. **No changes were made.**"
        )
    elif state == "no_pending":
        answer = (
            "Hiện không có thao tác Copilot nào đang chờ bạn xác nhận, nên mình **không thực hiện thay đổi nào**."
            if language == "vi" else
            "There is no pending Copilot action awaiting confirmation, so **no changes were made**."
        )
    else:
        answer = (
            "Mình không thể thực hiện thao tác đang chờ một cách an toàn. **Không có thay đổi nào được thực hiện.**"
            if language == "vi" else
            "I could not safely execute the pending action. **No changes were made.**"
        )
    return _result(message, answer, language, action or "none")


def _preview_progress_result(
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    route: RouteDecision,
    conversation_context: str = "",
) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "internship_progress",
            "Bạn chưa có kỳ thực tập trong hệ thống nên chưa thể chuẩn bị thao tác lưu progress.",
            "You do not have an internship record, so a progress-save action cannot be prepared.",
        )

    draft = _draft_progress_write(route, language)
    if not draft.has_work_details or not (draft.work_summary or "").strip():
        question = draft.clarification_question or (
            "Bạn muốn ghi nhận công việc/progress cụ thể nào?"
            if language == "vi" else
            "What specific work/progress would you like to record?"
        )
        return _clarification_result(message, language, "internship_progress", question)

    week = draft.week
    if week is None:
        # Week is optional for a progress log. If omitted, derive it from the
        # authenticated internship calendar rather than inventing it from wording.
        week = _parse_week_number(draft.work_summary or "", internship)
    hours = draft.hours

    if language == "vi":
        lines = [
            "**Mình định thực hiện thao tác sau:**",
            f"- **Kỳ thực tập:** {internship.get('position_title') or 'Internship'} tại {internship.get('company_name') or 'công ty hiện tại'}",
            f"- **Progress sẽ ghi:** {draft.work_summary}",
            f"- **Tuần:** {week if week is not None else 'không xác định'}",
            f"- **Số giờ sẽ cộng:** {f'{hours:g} giờ' if hours is not None else 'không cộng giờ (người dùng chưa cung cấp số giờ)'}",
            "- **Sau khi xác nhận:** tạo progress log; chỉ khi người dùng đã nêu số giờ thì mới cập nhật completed_hours/progress_percentage.",
            "",
            "**Chưa có dữ liệu nào được ghi.** Bạn có thể xác nhận để thực hiện, hủy, chỉnh sửa nội dung ở tin nhắn tiếp theo, hoặc tiếp tục hỏi việc khác.",
        ]
    else:
        lines = [
            "**I plan to perform this action:**",
            f"- **Internship:** {internship.get('position_title') or 'Internship'} at {internship.get('company_name') or 'current company'}",
            f"- **Progress to record:** {draft.work_summary}",
            f"- **Week:** {week if week is not None else 'not determined'}",
            f"- **Hours to add:** {f'{hours:g} hours' if hours is not None else 'no hours will be added (the user did not provide hours)'}",
            "- **After confirmation:** create a progress log; completed_hours/progress_percentage will be updated only when the user explicitly supplied hours.",
            "",
            "**Nothing has been written yet.** You can confirm, cancel, revise the draft in your next message, or continue with another question.",
        ]
    return _result(message, "\n".join(lines), language, "internship_progress")

def _preview_weekly_reflection_result(
    db: Session,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    route: RouteDecision,
    conversation_context: str = "",
) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "weekly_reflection",
            "Bạn chưa có kỳ thực tập trong hệ thống để chuẩn bị lưu reflection.",
            "You do not have an internship record, so a reflection-save action cannot be prepared.",
        )

    draft = _draft_reflection_write(route, language)
    if not draft.has_target_week or draft.week is None:
        question = draft.clarification_question or (
            "Bạn muốn lưu reflection của tuần nào?"
            if language == "vi" else
            "Which week would you like to save the reflection for?"
        )
        return _clarification_result(message, language, "weekly_reflection", question)

    week = int(draft.week)
    logs = _get_progress_logs(db, int(internship["id"]), week)
    existing = db.execute(
        text(
            """
            SELECT id, status
            FROM weekly_reports
            WHERE internship_id = :internship_id
              AND report_type = 'WEEKLY'
              AND week_number = :week
            LIMIT 1
            """
        ),
        {"internship_id": int(internship["id"]), "week": week},
    ).mappings().first()

    if not logs and not existing:
        return _clarification_result(
            message,
            language,
            "weekly_reflection",
            (
                f"Tuần {week} hiện chưa có progress/report để tạo reflection. Bạn muốn ghi progress cho tuần này trước không?"
                if language == "vi" else
                f"Week {week} has no progress/report data to build a reflection from. Would you like to record progress for that week first?"
            ),
        )

    if language == "vi":
        answer = (
            "**Mình định thực hiện thao tác sau:**\n"
            f"- Tổng hợp reflection **tuần {week}** từ dữ liệu progress/report đang có.\n"
            f"- Progress logs tìm thấy: **{len(logs)}**.\n"
            f"- Weekly report hiện có: **{'có' if existing else 'chưa có'}**.\n"
            "- Sau khi xác nhận: tạo/cập nhật weekly report và lưu nội dung ở trạng thái **DRAFT**.\n\n"
            "**Hiện chưa lưu hay sửa weekly report nào.** Bạn có thể xác nhận, hủy, chỉnh sửa bản nháp, hoặc tiếp tục hỏi việc khác."
        )
    else:
        answer = (
            "**I plan to perform this action:**\n"
            f"- Generate the **week {week}** reflection from stored progress/report data.\n"
            f"- Progress logs found: **{len(logs)}**.\n"
            f"- Existing weekly report: **{'yes' if existing else 'no'}**.\n"
            "- After confirmation: create/update the weekly report and save it as **DRAFT**.\n\n"
            "**No weekly report has been saved or changed yet.** You can confirm, cancel, revise the draft, or continue with another question."
        )
    return _result(message, answer, language, "weekly_reflection")

def _notification_preference_labels(key: str) -> tuple[str, str]:
    labels = {
        "report_deadline": ("nhắc deadline/report", "deadline/report notifications"),
        "lecturer_feedback": ("nhắc phản hồi giảng viên", "lecturer-feedback notifications"),
        "internship_status": ("thông báo trạng thái internship", "internship-status notifications"),
        "email_notifications": ("thông báo email", "email notifications"),
    }
    return labels.get(key, (key, key))


def _format_reminder_datetime(value: datetime, language: str) -> str:
    """Human-facing datetime; DB/API values remain timezone-aware datetimes."""
    if language == "vi":
        return value.strftime("%H:%M, %d/%m/%Y")
    return value.strftime("%H:%M, %d/%m/%Y")


def _resolve_reminder_schedule(
    db: Session,
    student_id: int,
    draft: ReminderWriteDraft,
    language: str,
) -> tuple[datetime | None, str | None, str | None, str | None]:
    """Resolve a validated semantic reminder draft into concrete schedule data.

    Returns (scheduled_at, title, detail, clarification_question). It fails closed
    when a deadline reference is ambiguous instead of silently picking the first.
    """
    now = _local_now()
    content = (draft.reminder_content or "").strip()
    scheduled: datetime | None = None
    title: str | None = None
    detail: str | None = None

    if draft.scheduled_at_iso:
        try:
            raw_iso = draft.scheduled_at_iso.strip()
            if raw_iso.endswith("Z"):
                raw_iso = raw_iso[:-1] + "+00:00"
            scheduled = datetime.fromisoformat(raw_iso)
            if scheduled.tzinfo is None and now.tzinfo is not None:
                scheduled = scheduled.replace(tzinfo=now.tzinfo)
        except (TypeError, ValueError):
            scheduled = None

        if scheduled is None:
            return None, None, None, (
                "Mình chưa thể xác nhận thời điểm nhắc đã được hiểu an toàn. Bạn cho mình một mốc thời gian khác nhé."
                if language == "vi" else
                "I couldn't safely validate the reminder time. Please provide another time."
            )
        title = f"Reminder: {content[:80]}" if content else "Internship reminder"
        detail = content

    elif draft.days_before is not None:
        upcoming = _get_upcoming_items(db, student_id)
        target = _choose_upcoming_target(draft.deadline_reference or "", upcoming)
        if target is None:
            return None, None, None, (
                "Mình chưa xác định được chính xác deadline nào bạn muốn dùng. Bạn hãy nói rõ tên deadline/report cần nhắc."
                if language == "vi" else
                "I could not determine exactly which deadline you mean. Please name the deadline/report to use."
            )
        due_at = target.get("due_at")
        if due_at is None:
            return None, None, None, (
                "Deadline này chưa có thời gian due_at trong hệ thống."
                if language == "vi" else
                "That deadline does not have a due_at time in the system."
            )
        if getattr(due_at, "tzinfo", None) is None and now.tzinfo is not None:
            due_at = due_at.replace(tzinfo=now.tzinfo)
        scheduled = due_at - timedelta(days=int(draft.days_before))
        title = f"Reminder: {target.get('title') or content or 'Internship deadline'}"
        detail = content or f"{target.get('title')} is due at {target.get('due_at')}"

    if scheduled is None:
        return None, None, None, (
            "Bạn muốn mình nhắc vào thời điểm nào?"
            if language == "vi" else
            "When would you like me to remind you?"
        )
    if scheduled <= now:
        return None, None, None, (
            "Thời điểm nhắc phải ở tương lai."
            if language == "vi" else
            "The reminder time must be in the future."
        )
    return scheduled, title, detail, None


def _preview_reminder_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    route: RouteDecision,
    conversation_context: str = "",
) -> QueryResult:
    draft = _draft_reminder_write(route, language)
    if not draft.has_required_details:
        question = draft.clarification_question or (
            "Bạn muốn mình nhắc việc gì và vào thời điểm nào?"
            if language == "vi" else
            "What would you like me to remind you about, and when?"
        )
        return _clarification_result(message, language, "smart_notifications", question)

    if draft.kind == "PREFERENCE":
        assert draft.preference_key is not None
        assert draft.preference_enabled is not None
        label_vi, label_en = _notification_preference_labels(draft.preference_key)
        state_vi = "bật" if draft.preference_enabled else "tắt"
        state_en = "enable" if draft.preference_enabled else "disable"
        answer = (
            f"**Mình định {state_vi} {label_vi}** cho tài khoản của bạn.\n\n"
            "**Hiện chưa thay đổi notification_preferences.** Bạn có thể xác nhận, hủy, chỉnh sửa yêu cầu, hoặc tiếp tục hỏi việc khác."
            if language == "vi" else
            f"**I plan to {state_en} {label_en}** for your account.\n\n"
            "**notification_preferences has not been changed yet.** You can confirm, cancel, revise the request, or continue with another question."
        )
        return _result(message, answer, language, "smart_notifications")

    scheduled, title, detail, question = _resolve_reminder_schedule(
        db, student_id, draft, language
    )
    if question or scheduled is None:
        return _clarification_result(
            message,
            language,
            "smart_notifications",
            question or (
                "Bạn muốn mình nhắc việc gì và vào thời điểm nào?"
                if language == "vi" else
                "What would you like me to remind you about, and when?"
            ),
        )

    if language == "vi":
        answer = (
            "**Mình định tạo reminder sau:**\n"
            f"- **Tiêu đề:** {title}\n"
            f"- **Nội dung:** {detail}\n"
            f"- **Thời điểm nhắc:** {_format_reminder_datetime(scheduled, language)}\n"
            "- **Sau khi xác nhận:** tạo một `calendar_events` reminder cho chính tài khoản của bạn.\n\n"
            "**Hiện chưa tạo reminder nào.** Bạn có thể xác nhận để tạo, hủy, đổi thời gian/nội dung ở tin nhắn tiếp theo, hoặc tiếp tục hỏi việc khác."
        )
    else:
        answer = (
            "**I plan to create this reminder:**\n"
            f"- **Title:** {title}\n"
            f"- **Content:** {detail}\n"
            f"- **Reminder time:** {_format_reminder_datetime(scheduled, language)}\n"
            "- **After confirmation:** create a `calendar_events` reminder for your account.\n\n"
            "**No reminder has been created yet.** You can confirm, cancel, change the time/content in your next message, or continue with another question."
        )
    return _result(message, answer, language, "smart_notifications")

def _preview_escalation_result(
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    route: RouteDecision,
    conversation_context: str = "",
) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "human_escalation",
            "Bạn chưa có kỳ thực tập trong hệ thống nên chưa thể chuẩn bị escalation gắn với internship.",
            "You do not have an internship record, so an internship escalation cannot be prepared.",
        )
    draft = _draft_escalation(route, language)
    if (
        not draft.has_incident_details
        or not (draft.subject or "").strip()
        or not (draft.description or "").strip()
    ):
        question = (draft.clarification_question or (
            "Bạn muốn escalation phản ánh vấn đề/sự việc gì trong kỳ thực tập?"
            if language == "vi" else
            "What internship issue or incident would you like the escalation to report?"
        )).strip()
        return _result(message, question, language, "human_escalation").model_copy(
            update={
                "guardrail_reason": "copilot_clarification_required",
                "groundedness_status": "skip",
                "groundedness_reason": "copilot_clarification_required",
            }
        )

    target = "Faculty Mentor" if draft.target == "FACULTY_MENTOR" else "CAID queue"
    lecturer_name = internship.get("lecturer_name")
    if language == "vi":
        delivery = (
            f"Sau khi xác nhận, hệ thống sẽ tạo escalation record và notification/cảnh báo cho Faculty Mentor **{lecturer_name}**."
            if draft.target == "FACULTY_MENTOR" and internship.get("lecturer_id") else
            "Sau khi xác nhận, hệ thống sẽ tạo escalation record. Hiện internship chưa có Faculty Mentor được gán nên chưa thể hứa sẽ gửi notification cho một người nhận cụ thể."
            if draft.target == "FACULTY_MENTOR" else
            "Sau khi xác nhận, hệ thống sẽ tạo escalation record với đích CAID queue; không tự nhận là đã gửi email nếu backend không có email integration."
        )
        answer = (
            "**Mình định tạo escalation với nội dung sau:**\n"
            f"- **Chủ đề:** {draft.subject or '-'}\n"
            f"- **Nội dung sẽ ghi nhận:** {draft.description or '-'}\n"
            f"- **Loại:** {draft.escalation_type}\n"
            f"- **Mức độ:** {draft.severity}\n"
            f"- **Nơi nhận:** {target}\n"
            f"- **Kỳ thực tập:** {internship.get('position_title') or 'Internship'} tại {internship.get('company_name') or 'công ty hiện tại'}\n"
            f"- **Khi thực hiện:** {delivery}\n\n"
            "**Hiện chưa tạo escalation, chưa gửi cảnh báo và chưa tạo notification.** Bạn có thể xác nhận, hủy, chỉnh sửa nội dung/người nhận/mức độ, hoặc tiếp tục hỏi việc khác."
        )
    else:
        delivery = (
            f"After confirmation, the system will create the escalation record and a warning/notification for assigned Faculty Mentor **{lecturer_name}**."
            if draft.target == "FACULTY_MENTOR" and internship.get("lecturer_id") else
            "After confirmation, the escalation record will be created. No Faculty Mentor is currently assigned, so no specific recipient notification can be promised."
            if draft.target == "FACULTY_MENTOR" else
            "After confirmation, the escalation record will be created for the CAID queue; the system will not claim an email was sent without an email integration."
        )
        answer = (
            "**I plan to create this escalation:**\n"
            f"- **Subject:** {draft.subject or '-'}\n"
            f"- **Content to record:** {draft.description or '-'}\n"
            f"- **Type:** {draft.escalation_type}\n"
            f"- **Severity:** {draft.severity}\n"
            f"- **Target:** {target}\n"
            f"- **Internship:** {internship.get('position_title') or 'Internship'} at {internship.get('company_name') or 'current company'}\n"
            f"- **When executed:** {delivery}\n\n"
            "**No escalation, warning, or notification has been created yet.** You can confirm, cancel, revise the content/recipient/severity, or continue with another question."
        )
    return _result(message, answer, language, "human_escalation")

# ---------------------------------------------------------------------------
# Common DB readers
# ---------------------------------------------------------------------------

def _get_student_profile(
    db: Session,
    student_id: int,
    *,
    requested_fields: set[str],
) -> dict[str, Any] | None:
    allowed = {
        "faculty": "sp.faculty",
        "major": "sp.major",
        "cohort": "sp.cohort",
        "gpa": "sp.gpa",
        "skills": "sp.skills",
    }
    selected = [field for field in requested_fields if field in allowed]
    if not selected:
        return None
    select_sql = ", ".join(f"{allowed[field]} AS {field}" for field in selected)
    row = db.execute(
        text(
            f"""
            SELECT {select_sql}
            FROM users AS u
            LEFT JOIN student_profiles AS sp ON sp.student_id = u.id
            WHERE u.id = :student_id AND u.role = 'STUDENT'
            LIMIT 1
            """
        ),
        {"student_id": student_id},
    ).mappings().first()
    return dict(row) if row else None


def _get_current_internship(db: Session, student_id: int) -> dict[str, Any] | None:
    row = db.execute(
        text(
            """
            SELECT i.id, i.position_title, i.description, i.start_date, i.end_date,
                   i.required_hours, i.completed_hours, i.progress_percentage, i.status,
                   i.semester_id, i.company_id, i.lecturer_id,
                   c.name AS company_name, c.industry AS company_industry,
                   s.name AS semester_name, s.semester_code,
                   s.registration_start_date, s.registration_end_date,
                   lecturer.full_name AS lecturer_name
            FROM internships AS i
            LEFT JOIN companies AS c ON c.id = i.company_id
            LEFT JOIN semesters AS s ON s.id = i.semester_id
            LEFT JOIN users AS lecturer ON lecturer.id = i.lecturer_id
            WHERE i.student_id = :student_id
            ORDER BY
                CASE WHEN i.status IN ('IN_PROGRESS','NOT_STARTED','PAUSED') THEN 0 ELSE 1 END,
                i.start_date DESC NULLS LAST, i.id DESC
            LIMIT 1
            """
        ),
        {"student_id": student_id},
    ).mappings().first()
    return dict(row) if row else None


def _get_current_application(db: Session, student_id: int) -> dict[str, Any] | None:
    row = db.execute(
        text(
            """
            SELECT ia.id, ia.position_title, ia.internship_type, ia.description,
                   ia.expected_start_date, ia.expected_end_date, ia.work_mode,
                   ia.credits, ia.status, ia.submitted_at, ia.reviewed_at,
                   ia.semester_id, ia.company_id, ia.assigned_lecturer_id,
                   c.name AS company_name, c.industry AS company_industry,
                   s.name AS semester_name, s.semester_code,
                   s.registration_start_date, s.registration_end_date
            FROM internship_applications AS ia
            LEFT JOIN companies AS c ON c.id = ia.company_id
            LEFT JOIN semesters AS s ON s.id = ia.semester_id
            WHERE ia.student_id = :student_id
            ORDER BY
                CASE WHEN ia.status IN ('DRAFT','SUBMITTED','UNDER_REVIEW','APPROVED') THEN 0 ELSE 1 END,
                ia.created_at DESC, ia.id DESC
            LIMIT 1
            """
        ),
        {"student_id": student_id},
    ).mappings().first()
    return dict(row) if row else None


def _table_exists(db: Session, table_name: str) -> bool:
    return bool(
        db.execute(
            text("SELECT to_regclass(:name) IS NOT NULL"),
            {"name": f"public.{table_name}"},
        ).scalar()
    )


def _get_upcoming_items(db: Session, student_id: int, limit: int = 20) -> list[dict[str, Any]]:
    rows = db.execute(
        text(
            """
            WITH current_internship AS (
                SELECT id, semester_id
                FROM internships
                WHERE student_id = :student_id
                ORDER BY
                    CASE WHEN status IN ('IN_PROGRESS','NOT_STARTED','PAUSED') THEN 0 ELSE 1 END,
                    start_date DESC NULLS LAST, id DESC
                LIMIT 1
            )
            SELECT * FROM (
                SELECT 'report' AS source_type, wr.id AS source_id,
                       COALESCE(wr.title, wr.report_type || ' report') AS title,
                       wr.due_at, wr.status, wr.report_type AS detail
                FROM weekly_reports AS wr
                JOIN current_internship AS ci ON ci.id = wr.internship_id
                WHERE wr.due_at IS NOT NULL
                  AND wr.due_at >= NOW()
                  AND wr.status IN ('DRAFT','REVISION_REQUIRED')

                UNION ALL

                SELECT 'checklist' AS source_type, ci_item.id AS source_id,
                       ci_item.title, ci_item.due_at, ci_item.status,
                       ci_item.category AS detail
                FROM checklist_items AS ci_item
                JOIN current_internship AS ci ON ci.id = ci_item.internship_id
                WHERE ci_item.due_at IS NOT NULL
                  AND ci_item.due_at >= NOW()
                  AND ci_item.status <> 'COMPLETED'

                UNION ALL

                SELECT 'deadline' AS source_type, d.id AS source_id,
                       d.title, d.due_at, 'ACTIVE' AS status,
                       d.deadline_type AS detail
                FROM deadlines AS d
                JOIN current_internship AS ci ON ci.semester_id = d.semester_id
                WHERE d.is_active = TRUE
                  AND d.due_at >= NOW()
                  AND (d.target_role IS NULL OR d.target_role IN ('STUDENT','ALL'))
            ) AS upcoming
            ORDER BY due_at ASC
            LIMIT :limit
            """
        ),
        {"student_id": student_id, "limit": limit},
    ).mappings().all()
    return [dict(row) for row in rows]


def _get_open_checklist(db: Session, internship_id: int) -> list[dict[str, Any]]:
    rows = db.execute(
        text(
            """
            SELECT id, title, description, category, priority, status, due_at
            FROM checklist_items
            WHERE internship_id = :internship_id AND status <> 'COMPLETED'
            ORDER BY
                CASE priority WHEN 'HIGH' THEN 0 WHEN 'MEDIUM' THEN 1 ELSE 2 END,
                due_at ASC NULLS LAST, id ASC
            """
        ),
        {"internship_id": internship_id},
    ).mappings().all()
    return [dict(row) for row in rows]


def _get_reports(db: Session, internship_id: int, limit: int = 20) -> list[dict[str, Any]]:
    rows = db.execute(
        text(
            """
            SELECT id, week_number, report_type, title, status, due_at,
                   submitted_at, lecturer_feedback, lecturer_score, content
            FROM weekly_reports
            WHERE internship_id = :internship_id
            ORDER BY COALESCE(due_at, created_at) DESC
            LIMIT :limit
            """
        ),
        {"internship_id": internship_id, "limit": limit},
    ).mappings().all()
    return [dict(row) for row in rows]


def _get_evaluations(db: Session, internship_id: int) -> list[dict[str, Any]]:
    rows = db.execute(
        text(
            """
            SELECT id, evaluator_type, evaluation_type, total_score, feedback,
                   strengths, improvements, status, submitted_at
            FROM evaluations
            WHERE internship_id = :internship_id
            ORDER BY created_at DESC
            """
        ),
        {"internship_id": internship_id},
    ).mappings().all()
    return [dict(row) for row in rows]


def _get_progress_logs(db: Session, internship_id: int, week_number: int | None = None) -> list[dict[str, Any]]:
    if not _table_exists(db, "internship_progress_logs"):
        return []
    rows = db.execute(
        text(
            """
            SELECT id, log_date, week_number, title, description, hours,
                   skills, blockers, created_at
            FROM internship_progress_logs
            WHERE internship_id = :internship_id
              AND (:week_number IS NULL OR week_number = :week_number)
            ORDER BY log_date DESC, created_at DESC
            LIMIT 100
            """
        ),
        {"internship_id": internship_id, "week_number": week_number},
    ).mappings().all()
    return [dict(row) for row in rows]


# ---------------------------------------------------------------------------
# RAG helpers
# ---------------------------------------------------------------------------

def _internship_route(intent: str, query: str, action: str, language: str) -> RouteDecision:
    return RouteDecision(
        intent=intent,
        scope="internship",
        language=language if language in {"vi", "en"} else "vi",
        allowed_document_types=list(INTERNSHIP_DOCUMENT_TYPES),
        blocked_document_types=[
            t for t in ALL_ROUTED_DOCUMENT_TYPES if t not in INTERNSHIP_DOCUMENT_TYPES
        ],
        retrieval_query=query,
        evidence_mode="semantic",
        assistant_action=action,
        action_mode="inform",
        reason="copilot_internal_policy_lookup",
    )


def _career_route(query: str, action: str, language: str) -> RouteDecision:
    return RouteDecision(
        intent="career_opportunity",
        scope="career",
        language=language if language in {"vi", "en"} else "vi",
        allowed_document_types=list(CAREER_DOCUMENT_TYPES),
        blocked_document_types=[
            t for t in ALL_ROUTED_DOCUMENT_TYPES if t not in CAREER_DOCUMENT_TYPES
        ],
        retrieval_query=query,
        evidence_mode="semantic",
        assistant_action=action,
        action_mode="inform",
        reason="copilot_internal_career_lookup",
    )


def _lookup_rag(
    rag_lookup: RagLookup | None,
    query: str,
    route: RouteDecision,
) -> QueryResult | None:
    if rag_lookup is None:
        return None
    try:
        return rag_lookup(query, route)
    except Exception as exc:  # pragma: no cover - network/service fallback
        logger.warning("Copilot RAG lookup failed: %s", exc)
        return None


# ---------------------------------------------------------------------------
# LLM helper — reuses the existing cached chat model/client.
# ---------------------------------------------------------------------------

def _llm_text(system_prompt: str, user_prompt: str, fallback: str) -> str:
    settings = get_settings()
    if not settings.openai_api_key:
        return fallback
    try:
        llm = _get_chat_llm(settings.openai_chat_model or settings.model_name, 0.2)
        response = llm.invoke(
            [("system", system_prompt), ("human", user_prompt)],
            config={"callbacks": langfuse_callbacks()},
        )
        content = getattr(response, "content", response)
        if isinstance(content, list):
            content = "".join(
                str(item.get("text", "")) if isinstance(item, dict) else str(item)
                for item in content
            )
        return str(content or "").strip() or fallback
    except Exception as exc:  # pragma: no cover
        logger.warning("Copilot LLM generation failed: %s", exc)
        return fallback


def _result(
    message: str,
    answer: str,
    language: str,
    action: str,
    *,
    route_scope: str = "personal_student",
    confidence: float = 1.0,
    sources: list[dict] | None = None,
) -> QueryResult:
    return QueryResult(
        query=message,
        answer=answer,
        answer_status="answered",
        answer_language=language,
        confidence=confidence,
        sources=sources or [],
        route_intent=f"copilot_{action}",
        route_scope=route_scope,
        guardrail_passed=True,
        groundedness_status="skip" if not sources else "supported",
        groundedness_reason="copilot_db_action" if not sources else "copilot_rag_plus_db",
    )


def _not_found(message: str, language: str, action: str, detail_vi: str, detail_en: str) -> QueryResult:
    return QueryResult(
        query=message,
        answer=detail_vi if language == "vi" else detail_en,
        answer_status="not_found",
        answer_language=language,
        confidence=0.0,
        sources=[],
        route_intent=f"copilot_{action}",
        route_scope="personal_student",
        guardrail_passed=True,
    )


# ---------------------------------------------------------------------------
# 1. Eligibility Checker
# ---------------------------------------------------------------------------

def _eligibility_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    rag_lookup: RagLookup | None,
) -> QueryResult:
    profile, internship, application = snapshot.profile, snapshot.internship, snapshot.application
    if not profile:
        return _not_found(
            message, language, "eligibility_checker",
            "Hệ thống chưa có hồ sơ sinh viên để kiểm tra điều kiện thực tập của bạn.",
            "No student profile is available to check your internship eligibility.",
        )

    policy_query = (
        "All baseline requirements for a valid credit-bearing internship: Internship Request Form approval, "
        "registration timing, Foundation courses, GPA, full-time and part-time duration/hours, evaluations, reflection, "
        "and internship readiness or orientation prerequisites."
    )
    policy_result = _lookup_rag(
        rag_lookup,
        policy_query,
        _internship_route("internship_eligibility", policy_query, "eligibility_checker", language),
    )

    facts = {
        "gpa": profile.get("gpa"),
        "major": profile.get("major"),
        "application": application,
        "internship": internship,
    }
    policy_text = policy_result.answer if policy_result else ""

    fallback = (
        "Mình đã lấy được dữ liệu hồ sơ của bạn nhưng chưa thể đối chiếu đầy đủ với tài liệu chính thức. "
        "Vui lòng thử lại khi RAG policy khả dụng."
        if language == "vi" else
        "I retrieved your profile, but the official policy evidence is not currently available for a complete eligibility check."
    )
    system = (
        "You are Internova's eligibility checker. Use ONLY the authenticated DB facts and the supplied official-policy answer. "
        "Do not invent missing values. Separate the conclusion into: Satisfied, Not satisfied, Unknown/needs evidence. "
        "A missing DB field must be Unknown, never assumed. State that the chatbot does not approve internships. "
        "Answer in Vietnamese." if language == "vi" else
        "You are Internova's eligibility checker. Use ONLY the authenticated DB facts and supplied official-policy answer. "
        "Do not invent missing values. Separate: Satisfied, Not satisfied, Unknown/needs evidence. "
        "Missing DB fields are Unknown. The chatbot does not approve internships. Answer in English."
    )
    answer = _llm_text(
        system,
        f"STUDENT FACTS:\n{json.dumps(facts, ensure_ascii=False, default=str)}\n\nOFFICIAL POLICY ANSWER:\n{policy_text}\n\nQUESTION:\n{message}",
        fallback,
    )
    return _result(
        message,
        answer,
        language,
        "eligibility_checker",
        route_scope="internship" if policy_result else "personal_student",
        confidence=min(1.0, policy_result.confidence if policy_result else 0.6),
        sources=(policy_result.sources if policy_result else []),
    )


# ---------------------------------------------------------------------------
# 2 / 19. Checklist + personalized dashboard
# ---------------------------------------------------------------------------

def _journey_payload(
    db: Session,
    student_id: int,
    snapshot: StudentSnapshot,
    *,
    allowed_sections: set[str],
) -> dict[str, Any]:
    internship = snapshot.internship
    application = snapshot.application
    payload: dict[str, Any] = {
        "application": application if "applications" in allowed_sections else None,
        "internship": internship if "internship" in allowed_sections else None,
        "open_checklist": [],
        "upcoming": _get_upcoming_items(db, student_id) if "deadlines" in allowed_sections else [],
        "reports": [],
        "evaluations": [],
        "documents": [],
    }
    if internship:
        iid = int(internship["id"])
        if "checklist" in allowed_sections:
            payload["open_checklist"] = _get_open_checklist(db, iid)
        if "reports" in allowed_sections:
            payload["reports"] = _get_reports(db, iid)
        if "evaluations" in allowed_sections:
            payload["evaluations"] = _get_evaluations(db, iid)
        if "documents" in allowed_sections:
            payload["documents"] = [
                dict(row)
                for row in db.execute(
                    text(
                        """
                        SELECT document_type, title, status, uploaded_at
                        FROM internship_documents
                        WHERE internship_id = :internship_id
                        ORDER BY uploaded_at DESC
                        """
                    ),
                    {"internship_id": iid},
                ).mappings().all()
            ]
    elif application and "documents" in allowed_sections:
        payload["documents"] = [
            dict(row)
            for row in db.execute(
                text(
                    """
                    SELECT document_type, title, created_at
                    FROM application_documents
                    WHERE application_id = :application_id
                    ORDER BY created_at DESC
                    """
                ),
                {"application_id": int(application["id"])},
            ).mappings().all()
        ]
    return payload


def _dashboard_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    allowed_sections: set[str],
) -> QueryResult:
    data = _journey_payload(
        db, student_id, snapshot, allowed_sections=allowed_sections
    )
    internship = snapshot.internship if "internship" in allowed_sections else None
    application = snapshot.application if "applications" in allowed_sections else None
    if not internship and not application:
        return _not_found(
            message, language, "personalized_dashboard",
            "Hiện hệ thống chưa ghi nhận hồ sơ đăng ký hoặc kỳ thực tập nào trong phạm vi bạn yêu cầu.",
            "No internship application or internship record is available within the requested scope.",
        )

    if language == "vi":
        lines = ["**Internship Dashboard**"]
        if internship:
            lines.append(
                f"- Kỳ thực tập: {internship.get('position_title') or 'Chưa có vị trí'}"
                f" tại {internship.get('company_name') or 'chưa có công ty'} — trạng thái **{internship.get('status')}**."
            )
            if internship.get("required_hours"):
                lines.append(
                    f"- Giờ thực tập: **{internship.get('completed_hours', 0)}/{internship.get('required_hours')} giờ**; "
                    f"tiến độ hệ thống: **{internship.get('progress_percentage', 0)}%**."
                )
        elif application:
            lines.append(
                f"- Hồ sơ đăng ký: {application.get('position_title') or 'chưa có vị trí'}"
                f" tại {application.get('company_name') or 'chưa có công ty'} — trạng thái **{application.get('status')}**."
            )

        open_items = data["open_checklist"]
        upcoming = data["upcoming"]
        pending_reports = [r for r in data["reports"] if r.get("status") in {"DRAFT", "REVISION_REQUIRED", "LATE"}]
        if "checklist" in allowed_sections:
            lines.append(f"- Checklist chưa hoàn thành: **{len(open_items)}**.")
        if "deadlines" in allowed_sections:
            lines.append(f"- Deadline sắp tới: **{len(upcoming)}**.")
        if "reports" in allowed_sections:
            lines.append(f"- Báo cáo cần chú ý: **{len(pending_reports)}**.")
        if upcoming:
            lines.append("\n**Ưu tiên gần nhất:**")
            for item in upcoming[:3]:
                lines.append(f"- {item['title']} — {item['due_at']}")
        if open_items:
            lines.append("\n**Việc nên làm tiếp:**")
            for item in open_items[:3]:
                lines.append(f"- {item['title']} ({item.get('priority', 'MEDIUM')})")
        answer = "\n".join(lines)
    else:
        lines = ["**Internship Dashboard**"]
        if internship:
            lines.append(
                f"- Internship: {internship.get('position_title') or 'No position'}"
                f" at {internship.get('company_name') or 'No company'} — **{internship.get('status')}**."
            )
            if internship.get("required_hours"):
                lines.append(
                    f"- Hours: **{internship.get('completed_hours', 0)}/{internship.get('required_hours')}**; "
                    f"system progress: **{internship.get('progress_percentage', 0)}%**."
                )
        elif application:
            lines.append(
                f"- Application: {application.get('position_title') or 'No position'}"
                f" at {application.get('company_name') or 'No company'} — **{application.get('status')}**."
            )

        open_items = data["open_checklist"]
        upcoming = data["upcoming"]
        pending_reports = [r for r in data["reports"] if r.get("status") in {"DRAFT", "REVISION_REQUIRED", "LATE"}]
        if "checklist" in allowed_sections:
            lines.append(f"- Open checklist items: **{len(open_items)}**.")
        if "deadlines" in allowed_sections:
            lines.append(f"- Upcoming deadlines: **{len(upcoming)}**.")
        if "reports" in allowed_sections:
            lines.append(f"- Reports needing attention: **{len(pending_reports)}**.")
        if upcoming:
            lines.append("\n**Nearest priorities:**")
            for item in upcoming[:3]:
                lines.append(f"- {item['title']} — {item['due_at']}")
        if open_items:
            lines.append("\n**Next actions:**")
            for item in open_items[:3]:
                lines.append(f"- {item['title']} ({item.get('priority', 'MEDIUM')})")
        answer = "\n".join(lines)
    return _result(message, answer, language, "personalized_dashboard")


def _checklist_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    rag_lookup: RagLookup | None,
    *,
    allowed_sections: set[str],
) -> QueryResult:
    data = _journey_payload(
        db, student_id, snapshot, allowed_sections=allowed_sections
    )
    policy_query = (
        "Internship approval and completion process: required steps before internship, during internship, reports, "
        "evaluations, reflection, and completion requirements."
    )
    policy_result = _lookup_rag(
        rag_lookup,
        policy_query,
        _internship_route("internship_registration", policy_query, "internship_checklist", language),
    )
    fallback = _dashboard_result(
        db, student_id, message, language, snapshot, allowed_sections=allowed_sections
    ).answer
    answer = _llm_text(
        (
            "Create a concise personalized internship checklist. Use stored facts only for completion status and the official-policy answer only for required steps. "
            "Separate Done, Next, and Unknown/needs confirmation. Never mark a policy requirement completed unless DB evidence supports it. Answer in Vietnamese."
            if language == "vi" else
            "Create a concise personalized internship checklist. Use stored facts only for completion status and the official-policy answer only for required steps. "
            "Separate Done, Next, and Unknown/needs confirmation. Never mark a policy requirement completed without DB evidence. Answer in English."
        ),
        f"STORED JOURNEY DATA:\n{json.dumps(data, ensure_ascii=False, default=str)}\n\nOFFICIAL POLICY:\n{policy_result.answer if policy_result else ''}\n\nQUESTION:\n{message}",
        fallback,
    )
    return _result(
        message,
        answer,
        language,
        "internship_checklist",
        route_scope="internship" if policy_result else "personal_student",
        confidence=(policy_result.confidence if policy_result else 0.8),
        sources=(policy_result.sources if policy_result else []),
    )


# ---------------------------------------------------------------------------
# 5. Deadline / Timeline
# ---------------------------------------------------------------------------

def _deadline_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
) -> QueryResult:
    items = _get_upcoming_items(db, student_id)
    if not items:
        return _not_found(
            message, language, "deadline_timeline",
            "Hiện hệ thống chưa ghi nhận deadline sắp tới nào cho kỳ thực tập của bạn.",
            "There are currently no upcoming internship deadlines recorded for you.",
        )
    title = "**Deadline sắp tới của bạn:**" if language == "vi" else "**Your upcoming deadlines:**"
    lines = [title]
    for item in items[:10]:
        lines.append(f"- {item['title']} — **{item['due_at']}** ({item['source_type']})")
    return _result(message, "\n".join(lines), language, "deadline_timeline")


# ---------------------------------------------------------------------------
# 6. Internship matching
# ---------------------------------------------------------------------------

def _normalize_skill(value: str) -> str:
    return re.sub(r"[^a-z0-9+#.]+", " ", value.lower()).strip()


def _extract_skills_from_cv(cv_text: str) -> list[str]:
    if not cv_text.strip():
        return []
    settings = get_settings()
    if not settings.openai_api_key:
        return []
    try:
        llm = _get_chat_llm(settings.openai_chat_model or settings.model_name, 0.0)
        structured = llm.with_structured_output(ExtractedSkills)
        result = structured.invoke(
            [
                (
                    "system",
                    "Extract only concrete skills explicitly evidenced in the supplied CV. "
                    "Return concise canonical skill names. Do not infer skills from job titles alone and do not invent anything.",
                ),
                ("human", cv_text[:14000]),
            ],
            config={"callbacks": langfuse_callbacks()},
        )
        parsed = ExtractedSkills.model_validate(result)
        return [str(skill).strip() for skill in parsed.skills if str(skill).strip()]
    except Exception as exc:  # pragma: no cover
        logger.warning("CV skill extraction failed: %s", exc)
        return []


def _matching_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    use_stored_cv: bool,
) -> QueryResult:
    profile = snapshot.profile or {}
    raw_skills = [str(value).strip() for value in (profile.get("skills") or []) if str(value).strip()]
    cv_text = ""
    if use_stored_cv:
        cv_text, _ = _stored_cv_text(db, student_id)
    if cv_text:
        for skill in _extract_skills_from_cv(cv_text):
            if skill.casefold() not in {item.casefold() for item in raw_skills}:
                raw_skills.append(skill)
    skills = [_normalize_skill(value) for value in raw_skills if _normalize_skill(value)]
    if not skills:
        return _not_found(
            message, language, "internship_matching",
            "Hồ sơ/CV của bạn chưa có kỹ năng đọc được để matching. Hãy cập nhật skills hoặc upload CV có nội dung text.",
            "No readable skills were found in your profile/CV for matching. Update your skills or upload a text-readable CV.",
        )
    if not _table_exists(db, "internship_opportunities"):
        return _not_found(
            message, language, "internship_matching",
            "Chưa có bảng `internship_opportunities`. Hãy chạy migration Internship Copilot trước khi dùng matching với vị trí đang mở.",
            "The `internship_opportunities` table is not available. Run the Internship Copilot migration before matching open roles.",
        )

    rows = db.execute(
        text(
            """
            SELECT io.id, io.title, io.department, io.description, io.requirements,
                   io.skills_required, io.eligible_majors, io.work_mode, io.min_gpa,
                   io.application_deadline, c.name AS company_name, c.industry
            FROM internship_opportunities AS io
            JOIN companies AS c ON c.id = io.company_id
            WHERE io.status = 'OPEN'
              AND (io.application_deadline IS NULL OR io.application_deadline >= CURRENT_DATE)
            ORDER BY io.application_deadline ASC NULLS LAST, io.created_at DESC
            LIMIT 100
            """
        )
    ).mappings().all()
    if not rows:
        return _not_found(
            message, language, "internship_matching",
            "Hiện chưa có vị trí internship đang mở trong hệ thống để matching.",
            "There are currently no open internship opportunities in the system to match against.",
        )

    profile_set = set(skills)
    scored: list[tuple[float, dict[str, Any], list[str], list[str]]] = []
    gpa = float(profile.get("gpa")) if profile.get("gpa") is not None else None
    major = str(profile.get("major") or "").strip().casefold() or None
    for raw in rows:
        row = dict(raw)
        required = [_normalize_skill(str(s)) for s in (row.get("skills_required") or []) if str(s).strip()]
        req_set = set(required)
        matched = sorted(req_set & profile_set)
        missing = sorted(req_set - profile_set)
        if req_set:
            skill_score = len(matched) / len(req_set)
        else:
            haystack = " ".join(
                str(row.get(k) or "") for k in ("title", "description", "requirements")
            ).lower()
            hits = [s for s in profile_set if s and s in haystack]
            skill_score = min(1.0, len(hits) / max(1, min(5, len(profile_set))))
            matched = sorted(hits)
        weighted_parts: list[tuple[float, float]] = [(0.75, skill_score)]

        eligible_majors = [
            str(value).strip().casefold()
            for value in (row.get("eligible_majors") or [])
            if str(value).strip()
        ]
        if eligible_majors:
            if major is None:
                major_score = 0.5  # unknown rather than an assumed mismatch
            else:
                major_score = 1.0 if major in eligible_majors else 0.0
            weighted_parts.append((0.10, major_score))

        if row.get("min_gpa") is not None:
            if gpa is None:
                gpa_score = 0.5  # unknown, not treated as pass/fail
            else:
                gpa_score = 1.0 if gpa >= float(row["min_gpa"]) else 0.0
            weighted_parts.append((0.15, gpa_score))

        total_weight = sum(weight for weight, _ in weighted_parts)
        score = round(
            (sum(weight * value for weight, value in weighted_parts) / total_weight) * 100,
            1,
        )
        scored.append((score, row, matched, missing))

    scored.sort(key=lambda item: item[0], reverse=True)
    top = scored[:5]
    if language == "vi":
        lines = ["**Top vị trí phù hợp từ dữ liệu đang mở:**"]
        for score, row, matched, missing in top:
            lines.append(
                f"- **{score}%** — {row['title']} @ {row['company_name']}"
                + (f" | Khớp: {', '.join(matched[:5])}" if matched else "")
                + (f" | Thiếu: {', '.join(missing[:4])}" if missing else "")
            )
        lines.append("\nĐiểm trên là **ước lượng matching của hệ thống**, không phải điểm tuyển dụng chính thức.")
    else:
        lines = ["**Top matches from currently open roles:**"]
        for score, row, matched, missing in top:
            lines.append(
                f"- **{score}%** — {row['title']} @ {row['company_name']}"
                + (f" | Matched: {', '.join(matched[:5])}" if matched else "")
                + (f" | Missing: {', '.join(missing[:4])}" if missing else "")
            )
        lines.append("\nThese are **system-estimated match scores**, not official employer scores.")
    return _result(message, "\n".join(lines), language, "internship_matching")


# ---------------------------------------------------------------------------
# 7 / 8 / 9 / 13. CV, JD, interview, skill gap
# ---------------------------------------------------------------------------

def _bytes(value: Any) -> bytes | None:
    if value is None:
        return None
    if isinstance(value, bytes):
        return value
    if isinstance(value, memoryview):
        return value.tobytes()
    try:
        return bytes(value)
    except Exception:
        return None


def _extract_binary_document(data: bytes, file_name: str | None, mime_type: str | None) -> str:
    name = (file_name or "").lower()
    mime = (mime_type or "").lower()
    try:
        if name.endswith(".pdf") or "pdf" in mime:
            import fitz
            document = fitz.open(stream=data, filetype="pdf")
            return "\n".join(page.get_text("text") for page in document).strip()
        if name.endswith(".docx") or "wordprocessingml" in mime:
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join(parts).strip()
        return data.decode("utf-8", errors="ignore").strip()
    except Exception as exc:
        logger.warning("Could not extract uploaded internship document %s: %s", file_name, exc)
        return ""


def _get_latest_application_document(
    db: Session,
    student_id: int,
    document_type: str,
) -> dict[str, Any] | None:
    row = db.execute(
        text(
            """
            SELECT ad.id, ad.document_type, ad.title, ad.original_file_name,
                   ad.mime_type, ad.file_data, ad.created_at
            FROM application_documents AS ad
            WHERE ad.student_id = :student_id AND ad.document_type = :document_type
            ORDER BY ad.created_at DESC, ad.id DESC
            LIMIT 1
            """
        ),
        {"student_id": student_id, "document_type": document_type},
    ).mappings().first()
    return dict(row) if row else None


def _stored_cv_text(db: Session, student_id: int) -> tuple[str, dict[str, Any] | None]:
    candidates: list[dict[str, Any]] = []
    application_doc = _get_latest_application_document(db, student_id, "CV")
    if application_doc:
        candidates.append(application_doc)

    internship_doc = db.execute(
        text(
            """
            SELECT id, document_type, title, original_file_name, mime_type,
                   file_data, uploaded_at AS created_at
            FROM internship_documents
            WHERE student_id = :student_id AND document_type = 'CV'
            ORDER BY uploaded_at DESC, id DESC
            LIMIT 1
            """
        ),
        {"student_id": student_id},
    ).mappings().first()
    if internship_doc:
        candidates.append(dict(internship_doc))

    if not candidates:
        return "", None

    candidates.sort(
        key=lambda row: row.get("created_at") or datetime.min,
        reverse=True,
    )
    doc = candidates[0]
    data = _bytes(doc.get("file_data"))
    return (
        _extract_binary_document(data, doc.get("original_file_name"), doc.get("mime_type")) if data else "",
        doc,
    )


def _stored_jd_text(
    db: Session,
    student_id: int,
    snapshot: StudentSnapshot,
    *,
    allow_documents: bool = True,
    allow_application_context: bool = True,
) -> tuple[str, dict[str, Any] | None]:
    if allow_documents:
        doc = _get_latest_application_document(db, student_id, "JOB_DESCRIPTION")
        if doc:
            data = _bytes(doc.get("file_data"))
            text_value = _extract_binary_document(data, doc.get("original_file_name"), doc.get("mime_type")) if data else ""
            if text_value:
                return text_value, doc
    app = snapshot.application if allow_application_context else None
    if app and app.get("description"):
        return str(app["description"]), app
    internship = snapshot.internship
    if internship and internship.get("description"):
        return str(internship["description"]), internship
    return "", None


def _cv_improvement_result(
    db: Session, student_id: int, message: str, language: str, snapshot: StudentSnapshot,
    *, allow_application_context: bool,
) -> QueryResult:
    cv, meta = _stored_cv_text(db, student_id)
    if not cv:
        return _not_found(
            message, language, "cv_improvement",
            "Mình chưa đọc được CV đã lưu trong hệ thống. Hãy upload/paste CV hoặc gắn CV vào application để mình review.",
            "I could not read a stored CV. Upload/paste a CV or attach it to an application for review.",
        )
    jd, _ = _stored_jd_text(
        db, student_id, snapshot,
        allow_documents=allow_application_context,
        allow_application_context=allow_application_context,
    ) if allow_application_context else ("", None)
    answer = _llm_text(
        (
            "Review the supplied student CV for internship readiness. Never invent experience, metrics, projects, or skills. "
            "Give prioritized issues, concrete rewrites, and role-specific improvements. If a JD is supplied, tailor against it. Answer in Vietnamese."
            if language == "vi" else
            "Review the supplied student CV for internship readiness. Never invent experience, metrics, projects, or skills. "
            "Give prioritized issues, concrete rewrites, and role-specific improvements. If a JD is supplied, tailor against it. Answer in English."
        ),
        f"CV:\n{cv[:14000]}\n\nJD/ROLE CONTEXT:\n{jd[:8000]}\n\nREQUEST:\n{message}",
        ("Mình đã đọc CV nhưng model review hiện không khả dụng." if language == "vi" else "I read the CV, but the review model is currently unavailable."),
    )
    return _result(message, answer, language, "cv_improvement", confidence=0.9)


def _jd_result(
    db: Session, student_id: int, message: str, language: str, snapshot: StudentSnapshot,
    *, allow_documents: bool, allow_application_context: bool,
) -> QueryResult:
    jd, _ = _stored_jd_text(
        db, student_id, snapshot,
        allow_documents=allow_documents,
        allow_application_context=allow_application_context,
    )
    if not jd:
        return _not_found(
            message, language, "jd_analyzer",
            "Mình chưa tìm thấy Job Description trong application/kỳ thực tập của bạn. Hãy upload hoặc paste JD để phân tích.",
            "I could not find a Job Description in your application/internship. Upload or paste the JD to analyze it.",
        )
    answer = _llm_text(
        (
            "Analyze only the supplied internship job description. Extract responsibilities, must-haves, nice-to-haves, technical skills, soft skills, and interview signals. Answer in Vietnamese."
            if language == "vi" else
            "Analyze only the supplied internship job description. Extract responsibilities, must-haves, nice-to-haves, technical skills, soft skills, and interview signals. Answer in English."
        ),
        f"JOB DESCRIPTION:\n{jd[:16000]}\n\nREQUEST:\n{message}",
        jd[:3000],
    )
    return _result(message, answer, language, "jd_analyzer", confidence=0.9)


def _skill_gap_result(
    db: Session, student_id: int, message: str, language: str, snapshot: StudentSnapshot,
    *, allow_documents: bool, allow_application_context: bool,
) -> QueryResult:
    cv, _ = _stored_cv_text(db, student_id) if allow_documents else ("", None)
    jd, _ = _stored_jd_text(
        db, student_id, snapshot,
        allow_documents=allow_documents,
        allow_application_context=allow_application_context,
    )
    profile = snapshot.profile or {}
    profile_skills = profile.get("skills") or []
    if not jd:
        return _not_found(
            message, language, "skill_gap_analysis",
            "Chưa có JD/role requirement để so skill gap. Hãy upload/paste JD hoặc chọn một opportunity.",
            "There is no JD/role requirement to compare against. Upload/paste a JD or select an opportunity.",
        )
    if not cv and not profile_skills:
        return _not_found(
            message, language, "skill_gap_analysis",
            "Chưa có CV hoặc skills trong hồ sơ để so với JD.",
            "There is no CV or profile skill list to compare with the JD.",
        )
    answer = _llm_text(
        (
            "Compare only evidenced student skills/CV against the supplied JD. Use Strong / Partial / Missing / Priority. Never infer a skill that is not evidenced. Answer in Vietnamese."
            if language == "vi" else
            "Compare only evidenced student skills/CV against the supplied JD. Use Strong / Partial / Missing / Priority. Never infer a skill that is not evidenced. Answer in English."
        ),
        f"PROFILE SKILLS: {profile_skills}\n\nCV:\n{cv[:12000]}\n\nJD:\n{jd[:12000]}\n\nREQUEST:\n{message}",
        ("Không đủ dữ liệu để tạo skill-gap analysis." if language == "vi" else "There is not enough data for a skill-gap analysis."),
    )
    return _result(message, answer, language, "skill_gap_analysis", confidence=0.9)


def _interview_result(
    db: Session, student_id: int, message: str, language: str, snapshot: StudentSnapshot,
    *, allow_documents: bool, allow_application_context: bool,
) -> QueryResult:
    cv, _ = _stored_cv_text(db, student_id) if allow_documents else ("", None)
    jd, _ = _stored_jd_text(
        db, student_id, snapshot,
        allow_documents=allow_documents,
        allow_application_context=allow_application_context,
    )
    if not jd:
        return _not_found(
            message, language, "interview_preparation",
            "Chưa có JD/role trong hệ thống để tạo bộ phỏng vấn cá nhân hóa. Hãy upload/paste JD trước.",
            "No stored JD/role is available for personalized interview preparation. Upload or paste the JD first.",
        )
    answer = _llm_text(
        (
            "Prepare a realistic internship interview from the supplied JD and CV. Include technical, behavioral, and CV-verification questions; flag likely weak areas. Do not invent CV facts. Answer in Vietnamese."
            if language == "vi" else
            "Prepare a realistic internship interview from the supplied JD and CV. Include technical, behavioral, and CV-verification questions; flag likely weak areas. Do not invent CV facts. Answer in English."
        ),
        f"CV:\n{cv[:10000]}\n\nJD:\n{jd[:12000]}\n\nREQUEST:\n{message}",
        ("Mình chưa thể tạo bộ phỏng vấn lúc này." if language == "vi" else "I cannot generate the interview set right now."),
    )
    return _result(message, answer, language, "interview_preparation", confidence=0.9)


# ---------------------------------------------------------------------------
# 10 / 11. Progress + weekly reflection
# ---------------------------------------------------------------------------

def _parse_week_number(message: str, internship: dict[str, Any] | None) -> int | None:
    match = re.search(r"(?:tu[aầ]n|week)\s*#?\s*(\d{1,2})", message, re.IGNORECASE)
    if match:
        return max(1, int(match.group(1)))
    if internship and internship.get("start_date"):
        start = internship["start_date"]
        if isinstance(start, datetime):
            start = start.date()
        if isinstance(start, date):
            return max(1, ((date.today() - start).days // 7) + 1)
    return None


def _parse_hours(message: str) -> float | None:
    match = re.search(r"(\d+(?:[\.,]\d+)?)\s*(?:gi[oờ]|hours?|hrs?|h)\b", message, re.IGNORECASE)
    if not match:
        return None
    return float(match.group(1).replace(",", "."))


def _save_progress_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    route: RouteDecision,
    conversation_context: str = "",
) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "internship_progress",
            "Bạn chưa có kỳ thực tập trong hệ thống nên mình chưa thể lưu progress.",
            "You do not have an internship record, so progress cannot be saved yet.",
        )
    if not _table_exists(db, "internship_progress_logs"):
        return _not_found(
            message, language, "internship_progress",
            "Chưa có bảng progress log. Hãy chạy migration Internship Copilot trước.",
            "The progress-log table is missing. Run the Internship Copilot migration first.",
        )

    draft = _draft_progress_write(route, language)
    if not draft.has_work_details or not (draft.work_summary or "").strip():
        return _clarification_result(
            message,
            language,
            "internship_progress",
            draft.clarification_question or (
                "Bạn muốn ghi nhận công việc/progress cụ thể nào?"
                if language == "vi" else
                "What specific work/progress would you like to record?"
            ),
        )

    work_summary = " ".join((draft.work_summary or "").strip().split())
    week_number = draft.week
    if week_number is None:
        week_number = _parse_week_number(work_summary, internship)
    hours = draft.hours
    today = date.today()

    canonical_payload = json.dumps(
        {
            "work_summary": work_summary,
            "week": week_number,
            "hours": hours,
        },
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    )
    source_hash = hashlib.sha256(
        f"{student_id}:{internship['id']}:{today.isoformat()}:{canonical_payload}".encode("utf-8")
    ).hexdigest()

    inserted = db.execute(
        text(
            """
            INSERT INTO internship_progress_logs
                (internship_id, student_id, log_date, week_number, title, description, hours, source_hash)
            VALUES
                (:internship_id, :student_id, :log_date, :week_number, :title, :description, :hours, :source_hash)
            ON CONFLICT (student_id, internship_id, log_date, source_hash) DO NOTHING
            RETURNING id
            """
        ),
        {
            "internship_id": int(internship["id"]),
            "student_id": student_id,
            "log_date": today,
            "week_number": week_number,
            "title": f"Progress {today.isoformat()}",
            "description": work_summary,
            "hours": hours,
            "source_hash": source_hash,
        },
    ).mappings().first()

    if inserted is None:
        db.rollback()
        answer = (
            "Progress này đã được ghi nhận trước đó nên hệ thống không cộng giờ lần nữa."
            if language == "vi" else
            "This progress entry was already recorded, so the hours were not added again."
        )
        return _result(message, answer, language, "internship_progress")

    if hours is not None and hours > 0:
        aggregate_hours = int(round(hours))
        db.execute(
            text(
                """
                UPDATE internships
                SET completed_hours = GREATEST(0, completed_hours + :hours),
                    progress_percentage = CASE
                        WHEN required_hours IS NOT NULL AND required_hours > 0
                        THEN LEAST(100, ROUND(((completed_hours + :hours)::numeric / required_hours) * 100, 2))
                        ELSE progress_percentage
                    END,
                    updated_at = NOW()
                WHERE id = :internship_id AND student_id = :student_id
                """
            ),
            {"hours": aggregate_hours, "internship_id": int(internship["id"]), "student_id": student_id},
        )
    db.commit()

    if language == "vi":
        answer = (
            "Đã lưu progress với nội dung:\n"
            f"- **Công việc:** {work_summary}\n"
            f"- **Tuần:** {week_number if week_number is not None else 'không xác định'}\n"
            f"- **Số giờ:** {f'{hours:g} giờ' if hours is not None else 'không cộng giờ'}"
        )
    else:
        answer = (
            "Progress saved with:\n"
            f"- **Work:** {work_summary}\n"
            f"- **Week:** {week_number if week_number is not None else 'not determined'}\n"
            f"- **Hours:** {f'{hours:g} hours' if hours is not None else 'no hours added'}"
        )
    return _result(message, answer, language, "internship_progress")

def _progress_summary_result(db: Session, student_id: int, message: str, language: str, snapshot: StudentSnapshot) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "internship_progress",
            "Bạn chưa có kỳ thực tập trong hệ thống.",
            "You do not have an internship record in the system.",
        )
    logs = _get_progress_logs(db, int(internship["id"]))
    if not logs:
        return _not_found(
            message, language, "internship_progress",
            "Chưa có progress log nào được lưu cho kỳ thực tập này.",
            "No progress logs have been saved for this internship yet.",
        )
    answer = _llm_text(
        (
            "Summarize only the supplied progress-log records into completed work, hours, skills explicitly mentioned, blockers, and next actions. Answer in Vietnamese."
            if language == "vi" else
            "Summarize only the supplied progress-log records into completed work, hours, explicitly mentioned skills, blockers, and next actions. Answer in English."
        ),
        f"INTERNSHIP: {json.dumps(internship, ensure_ascii=False, default=str)}\n\nPROGRESS LOGS:\n{json.dumps(logs[:30], ensure_ascii=False, default=str)}\n\nREQUEST:\n{message}",
        "\n".join(f"- {x.get('log_date')}: {x.get('description')}" for x in logs[:10]),
    )
    return _result(message, answer, language, "internship_progress")


def _weekly_reflection_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    save_draft: bool,
    conversation_context: str = "",
    route: RouteDecision | None = None,
) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "weekly_reflection",
            "Bạn chưa có kỳ thực tập trong hệ thống để tạo reflection.",
            "You do not have an internship record to generate a reflection from.",
        )
    if save_draft:
        draft = _draft_reflection_write(route, language) if route is not None else ReflectionWriteDraft(
            has_target_week=False,
            clarification_question=("Bạn muốn lưu reflection của tuần nào?" if language == "vi" else "Which week would you like to save the reflection for?"),
        )
        if not draft.has_target_week or draft.week is None:
            return _clarification_result(
                message,
                language,
                "weekly_reflection",
                draft.clarification_question or (
                    "Bạn muốn lưu reflection của tuần nào?"
                    if language == "vi" else
                    "Which week would you like to save the reflection for?"
                ),
            )
        week = int(draft.week)
    else:
        week = _parse_week_number(message, internship)
        if week is None:
            return _not_found(
                message, language, "weekly_reflection",
                "Mình chưa xác định được tuần cần tổng hợp. Hãy nói rõ, ví dụ: `tuần 4`.",
                "I could not determine which week to summarize. Specify it, e.g. `week 4`.",
            )
    logs = _get_progress_logs(db, int(internship["id"]), week)
    existing = db.execute(
        text(
            """
            SELECT id, content, status, lecturer_feedback
            FROM weekly_reports
            WHERE internship_id = :internship_id
              AND report_type = 'WEEKLY'
              AND week_number = :week
            LIMIT 1
            """
        ),
        {"internship_id": int(internship["id"]), "week": week},
    ).mappings().first()
    if not logs and not (existing and existing.get("content")):
        return _not_found(
            message, language, "weekly_reflection",
            f"Chưa có progress/report content cho tuần {week}; mình không muốn bịa nội dung reflection.",
            f"There is no progress/report content for week {week}; I will not invent reflection content.",
        )
    answer = _llm_text(
        (
            "Write a concise internship weekly reflection using ONLY the supplied records. Include Completed work, Learning, Challenges, How handled, and Next-week goals. Do not invent tasks or outcomes. Answer in Vietnamese."
            if language == "vi" else
            "Write a concise internship weekly reflection using ONLY the supplied records. Include Completed work, Learning, Challenges, How handled, and Next-week goals. Do not invent tasks or outcomes. Answer in English."
        ),
        f"WEEK: {week}\nPROGRESS LOGS:\n{json.dumps(logs, ensure_ascii=False, default=str)}\n\nEXISTING REPORT:\n{json.dumps(dict(existing) if existing else {}, ensure_ascii=False, default=str)}\n\nREQUEST:\n{message}",
        (str(existing.get("content")) if existing and existing.get("content") else ""),
    )
    if save_draft:
        schedule = db.execute(
            text(
                """
                SELECT id, due_at FROM weekly_report_schedules
                WHERE semester_id = :semester_id AND week_number = :week
                LIMIT 1
                """
            ),
            {"semester_id": internship.get("semester_id"), "week": week},
        ).mappings().first() if internship.get("semester_id") else None
        if existing:
            db.execute(
                text(
                    """
                    UPDATE weekly_reports
                    SET content = :content, title = COALESCE(title, :title),
                        schedule_id = COALESCE(schedule_id, :schedule_id),
                        due_at = COALESCE(due_at, :due_at), updated_at = NOW()
                    WHERE id = :id
                    """
                ),
                {
                    "content": answer,
                    "title": f"Weekly Reflection - Week {week}",
                    "schedule_id": schedule.get("id") if schedule else None,
                    "due_at": schedule.get("due_at") if schedule else None,
                    "id": int(existing["id"]),
                },
            )
        else:
            db.execute(
                text(
                    """
                    INSERT INTO weekly_reports
                        (internship_id, schedule_id, week_number, report_type, title, content, status, due_at)
                    VALUES
                        (:internship_id, :schedule_id, :week, 'WEEKLY', :title, :content, 'DRAFT', :due_at)
                    """
                ),
                {
                    "internship_id": int(internship["id"]),
                    "schedule_id": schedule.get("id") if schedule else None,
                    "week": week,
                    "title": f"Weekly Reflection - Week {week}",
                    "content": answer,
                    "due_at": schedule.get("due_at") if schedule else None,
                },
            )
        db.commit()
        suffix = "\n\n**Đã lưu thành DRAFT trong weekly report.**" if language == "vi" else "\n\n**Saved as a DRAFT weekly report.**"
        answer += suffix
    return _result(message, answer, language, "weekly_reflection")


# ---------------------------------------------------------------------------
# 12. Evaluation preparation
# ---------------------------------------------------------------------------

def _evaluation_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    rag_lookup: RagLookup | None,
) -> QueryResult:
    internship = snapshot.internship
    evaluations = _get_evaluations(db, int(internship["id"])) if internship else []
    progress = _get_progress_logs(db, int(internship["id"])) if internship else []
    policy_query = "Internship evaluation criteria in Form 4 for faculty mentor and employer evaluation of intern."
    policy_result = _lookup_rag(
        rag_lookup,
        policy_query,
        _internship_route("internship_evaluation", policy_query, "evaluation_preparation", language),
    )
    if not policy_result and not evaluations:
        return _not_found(
            message, language, "evaluation_preparation",
            "Hiện chưa lấy được tiêu chí đánh giá chính thức hoặc evaluation record của bạn.",
            "Official evaluation criteria and your evaluation records are currently unavailable.",
        )
    answer = _llm_text(
        (
            "Prepare the student for internship evaluation. Official criteria may come only from the supplied policy/Form answer. Personal strengths/gaps may come only from the supplied evaluation/progress records. Do not invent ratings. Answer in Vietnamese."
            if language == "vi" else
            "Prepare the student for internship evaluation. Official criteria may come only from the supplied policy/Form answer. Personal strengths/gaps may come only from the supplied evaluation/progress records. Do not invent ratings. Answer in English."
        ),
        f"OFFICIAL CRITERIA:\n{policy_result.answer if policy_result else ''}\n\nEVALUATIONS:\n{json.dumps(evaluations, ensure_ascii=False, default=str)}\n\nPROGRESS:\n{json.dumps(progress[:30], ensure_ascii=False, default=str)}\n\nREQUEST:\n{message}",
        policy_result.answer if policy_result else "",
    )
    return _result(
        message,
        answer,
        language,
        "evaluation_preparation",
        route_scope="internship" if policy_result else "personal_student",
        confidence=(policy_result.confidence if policy_result else 0.8),
        sources=(policy_result.sources if policy_result else []),
    )


# ---------------------------------------------------------------------------
# 14. Career recommendation
# ---------------------------------------------------------------------------

def _career_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    rag_lookup: RagLookup | None,
) -> QueryResult:
    profile = snapshot.profile or {}
    query = "Career pathways and core competencies for VinUniversity colleges, especially the student's field."
    career_result = _lookup_rag(rag_lookup, query, _career_route(query, "career_recommendation", language))
    answer = _llm_text(
        (
            "Recommend career directions using ONLY the student's stored profile and the supplied Talent Handbook answer. Distinguish handbook facts from personalized recommendations. Answer in Vietnamese."
            if language == "vi" else
            "Recommend career directions using ONLY the student's stored profile and supplied Talent Handbook answer. Distinguish handbook facts from personalized recommendations. Answer in English."
        ),
        f"PROFILE:\n{json.dumps(profile, ensure_ascii=False, default=str)}\n\nTALENT HANDBOOK:\n{career_result.answer if career_result else ''}\n\nREQUEST:\n{message}",
        ("Chưa đủ dữ liệu để đưa ra định hướng cá nhân hóa." if language == "vi" else "There is not enough data for a personalized career recommendation."),
    )
    return _result(
        message,
        answer,
        language,
        "career_recommendation",
        route_scope="career" if career_result else "personal_student",
        confidence=(career_result.confidence if career_result else 0.7),
        sources=(career_result.sources if career_result else []),
    )


# ---------------------------------------------------------------------------
# 18. Smart notifications
# ---------------------------------------------------------------------------

def _local_now() -> datetime:
    settings = get_settings()
    tz_name = getattr(settings, "copilot_timezone", "Asia/Ho_Chi_Minh")
    try:
        return datetime.now(ZoneInfo(tz_name))
    except Exception:
        return datetime.now()


def _choose_upcoming_target(
    reference: str,
    items: list[dict[str, Any]],
) -> dict[str, Any] | None:
    """Resolve semantic deadline reference against DB entities without intent heuristics."""
    if not items:
        return None
    if len(items) == 1:
        return items[0]

    ref = " ".join((reference or "").strip().casefold().split())
    if not ref:
        return None

    exact: list[dict[str, Any]] = []
    contained: list[dict[str, Any]] = []
    for item in items:
        title = " ".join(str(item.get("title") or "").strip().casefold().split())
        if not title:
            continue
        if ref == title:
            exact.append(item)
        elif ref in title or title in ref:
            contained.append(item)

    if len(exact) == 1:
        return exact[0]
    if len(exact) > 1:
        return None
    if len(contained) == 1:
        return contained[0]
    return None


def _notification_preference_update_from_draft(
    db: Session,
    student_id: int,
    draft: ReminderWriteDraft,
    message: str,
    language: str,
) -> QueryResult:
    if draft.preference_key is None or draft.preference_enabled is None:
        return _clarification_result(
            message,
            language,
            "smart_notifications",
            (
                "Bạn muốn bật/tắt loại thông báo nào?"
                if language == "vi" else
                "Which notification category would you like to enable or disable?"
            ),
        )

    column = str(draft.preference_key)
    value = bool(draft.preference_enabled)
    label_vi, label_en = _notification_preference_labels(column)

    db.execute(
        text(
            f"""
            INSERT INTO notification_preferences (user_id, {column})
            VALUES (:user_id, :value)
            ON CONFLICT (user_id) DO UPDATE
            SET {column} = EXCLUDED.{column}, updated_at = NOW()
            """
        ),
        {"user_id": student_id, "value": value},
    )
    db.commit()
    state_vi = "bật" if value else "tắt"
    state_en = "enabled" if value else "disabled"
    return _result(
        message,
        f"Đã {state_vi} {label_vi}." if language == "vi" else f"{label_en.capitalize()} {state_en}.",
        language,
        "smart_notifications",
    )


def _schedule_reminder_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    route: RouteDecision,
    conversation_context: str = "",
) -> QueryResult:
    draft = _draft_reminder_write(route, language)
    if not draft.has_required_details:
        return _clarification_result(
            message,
            language,
            "smart_notifications",
            draft.clarification_question or (
                "Bạn muốn mình nhắc việc gì và vào thời điểm nào?"
                if language == "vi" else
                "What would you like me to remind you about, and when?"
            ),
        )

    if draft.kind == "PREFERENCE":
        return _notification_preference_update_from_draft(
            db, student_id, draft, message, language
        )

    scheduled, title, detail, question = _resolve_reminder_schedule(
        db, student_id, draft, language
    )
    if question or scheduled is None or title is None or detail is None:
        return _clarification_result(
            message,
            language,
            "smart_notifications",
            question or (
                "Bạn muốn mình nhắc việc gì và vào thời điểm nào?"
                if language == "vi" else
                "What would you like me to remind you about, and when?"
            ),
        )

    try:
        reminder_id = schedule_reminder(
            db=db,
            student_id=student_id,
            title=title,
            message=detail,
            scheduled_at=scheduled,
        )
    except ValueError as exc:
        return _not_found(
            message,
            language,
            "smart_notifications",
            str(exc),
            str(exc),
        )

    answer = (
        f"Đã tạo reminder **{title}** vào **{_format_reminder_datetime(scheduled, language)}** "
        f"(calendar event #{reminder_id}).\n- **Nội dung:** {detail}"
        if language == "vi" else
        f"Reminder **{title}** scheduled for **{_format_reminder_datetime(scheduled, language)}** "
        f"(calendar event #{reminder_id}).\n- **Content:** {detail}"
    )
    return _result(
        message,
        answer,
        language,
        "smart_notifications",
    )

def _reminder_overview_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    *,
    include_deadlines: bool,
) -> QueryResult:
    reminders = [
        dict(row)
        for row in get_pending_reminders(
            db,
            student_id,
            20,
        )
    ]

    upcoming = (
        _get_upcoming_items(db, student_id)
        if include_deadlines
        else []
    )

    if not reminders and not upcoming:
        return _not_found(
            message,
            language,
            "smart_notifications",
            "Bạn chưa có reminder pending hoặc deadline sắp tới trong hệ thống.",
            "You have no pending reminders or upcoming deadlines in the system.",
        )

    if language == "vi":
        lines = ["**Reminder đang chờ:**"]
        lines += (
            [
                f"- {row['title']} — {_format_reminder_datetime(row['scheduled_at'], language) if isinstance(row.get('scheduled_at'), datetime) else row.get('scheduled_at', '-')}"
                for row in reminders[:10]
            ]
            or ["- Chưa có reminder đã lên lịch."]
        )

        lines.append("\n**Deadline có thể cần nhắc:**")
        lines += (
            [
                f"- {item['title']} — {item['due_at']}"
                for item in upcoming[:5]
            ]
            or ["- Không có deadline sắp tới."]
        )
    else:
        lines = ["**Pending reminders:**"]
        lines += (
            [
                f"- {row['title']} — {_format_reminder_datetime(row['scheduled_at'], language) if isinstance(row.get('scheduled_at'), datetime) else row.get('scheduled_at', '-')}"
                for row in reminders[:10]
            ]
            or ["- No scheduled reminders."]
        )

        lines.append("\n**Upcoming deadlines worth reminding:**")
        lines += (
            [
                f"- {item['title']} — {item['due_at']}"
                for item in upcoming[:5]
            ]
            or ["- No upcoming deadlines."]
        )

    return _result(
        message,
        "\n".join(lines),
        language,
        "smart_notifications",
    )


# ---------------------------------------------------------------------------
# 20 / grievance execution. Human escalation is explicit-write only.
# ---------------------------------------------------------------------------


def _escalation_overview_result(db: Session, student_id: int, message: str, language: str) -> QueryResult:
    if not _table_exists(db, "internship_escalations"):
        return _not_found(
            message, language, "human_escalation",
            "Chưa có bảng escalation. Hãy chạy migration Internship Copilot trước.",
            "The escalation table is missing. Run the Internship Copilot migration first.",
        )

    rows = [
        dict(row)
        for row in db.execute(
            text(
                """
                SELECT
                    id,
                    internship_id,
                    lecturer_id,
                    escalation_type,
                    severity,
                    target,
                    subject,
                    description,
                    status,
                    created_at,
                    acknowledged_at,
                    resolved_at
                FROM internship_escalations
                WHERE student_id = :student_id
                ORDER BY created_at DESC
                LIMIT 50
                """
            ),
            {"student_id": student_id},
        ).mappings().all()
    ]

    if not rows:
        return _not_found(
            message, language, "human_escalation",
            "Bạn chưa có escalation nào trong hệ thống.",
            "You do not have any escalations in the system.",
        )

    if language == "vi":
        lines = ["**Escalation của bạn:**"]
        for row in rows[:10]:
            target = (
                "Faculty Mentor"
                if row["target"] == "FACULTY_MENTOR"
                else "CAID queue"
                if row["target"] == "CAID_QUEUE"
                else str(row["target"])
            )
            lines.extend(
                [
                    "",
                    f"### Escalation #{row['id']}",
                    f"- **Chủ đề:** {row['subject']}",
                    f"- **Nội dung đã ghi nhận:** {row['description']}",
                    f"- **Loại:** {row['escalation_type']}",
                    f"- **Mức độ:** {row['severity']}",
                    f"- **Nơi nhận:** {target}",
                    f"- **Trạng thái:** {row['status']}",
                    f"- **Thời điểm tạo:** {(row['created_at'].isoformat() if row['created_at'] is not None else '-')}",
                ]
            )
    else:
        lines = ["**Your escalations:**"]
        for row in rows[:10]:
            target = (
                "Faculty Mentor"
                if row["target"] == "FACULTY_MENTOR"
                else "CAID queue"
                if row["target"] == "CAID_QUEUE"
                else str(row["target"])
            )
            lines.extend(
                [
                    "",
                    f"### Escalation #{row['id']}",
                    f"- **Subject:** {row['subject']}",
                    f"- **Recorded content:** {row['description']}",
                    f"- **Type:** {row['escalation_type']}",
                    f"- **Severity:** {row['severity']}",
                    f"- **Target:** {target}",
                    f"- **Status:** {row['status']}",
                    f"- **Created at:** {(row['created_at'].isoformat() if row['created_at'] is not None else '-')}",
                ]
            )

    return _result(message, "\n".join(lines), language, "human_escalation")


def _create_escalation_result(
    db: Session,
    student_id: int,
    message: str,
    language: str,
    snapshot: StudentSnapshot,
    conversation_context: str = "",
    *,
    route: RouteDecision,
) -> QueryResult:
    internship = snapshot.internship
    if not internship:
        return _not_found(
            message, language, "human_escalation",
            "Bạn chưa có kỳ thực tập trong hệ thống nên chưa thể tạo escalation gắn với internship.",
            "You do not have an internship record, so an internship escalation cannot be created.",
        )
    if not _table_exists(db, "internship_escalations"):
        return _not_found(
            message, language, "human_escalation",
            "Chưa có bảng escalation. Hãy chạy migration Internship Copilot trước.",
            "The escalation table is missing. Run the Internship Copilot migration first.",
        )

    draft = _draft_escalation(route, language)
    if (
        not draft.has_incident_details
        or not (draft.subject or "").strip()
        or not (draft.description or "").strip()
    ):
        question = (draft.clarification_question or (
            "Bạn muốn escalation phản ánh vấn đề/sự việc gì trong kỳ thực tập?"
            if language == "vi" else
            "What internship issue or incident would you like the escalation to report?"
        )).strip()
        return _result(message, question, language, "human_escalation")

    lecturer_id = internship.get("lecturer_id")
    duplicate = db.execute(
        text(
            """
            SELECT id, severity, target, subject, description, status, created_at
            FROM internship_escalations
            WHERE student_id = :student_id
              AND internship_id = :internship_id
              AND description = :description
              AND created_at >= NOW() - INTERVAL '5 minutes'
            ORDER BY created_at DESC
            LIMIT 1
            """
        ),
        {
            "student_id": student_id,
            "internship_id": int(internship["id"]),
            "description": str(draft.description),
        },
    ).mappings().first()
    if duplicate:
        target = (
            "Faculty Mentor"
            if duplicate["target"] == "FACULTY_MENTOR"
            else "CAID queue"
            if duplicate["target"] == "CAID_QUEUE"
            else str(duplicate["target"])
        )
        if language == "vi":
            answer = (
                f"Escalation **#{duplicate['id']}** vừa được tạo trước đó nên hệ thống không tạo bản trùng.\n\n"
                f"- **Chủ đề:** {duplicate['subject']}\n"
                f"- **Nội dung đã ghi nhận:** {duplicate['description']}\n"
                f"- **Mức độ:** {duplicate['severity']}\n"
                f"- **Nơi nhận:** {target}\n"
                f"- **Trạng thái:** {duplicate['status']}\n"
                f"- **Thời điểm tạo:** {(duplicate['created_at'].isoformat() if duplicate['created_at'] is not None else '-')}"
            )
        else:
            answer = (
                f"Escalation **#{duplicate['id']}** was just created, so a duplicate was not created.\n\n"
                f"- **Subject:** {duplicate['subject']}\n"
                f"- **Recorded content:** {duplicate['description']}\n"
                f"- **Severity:** {duplicate['severity']}\n"
                f"- **Target:** {target}\n"
                f"- **Status:** {duplicate['status']}\n"
                f"- **Created at:** {(duplicate['created_at'].isoformat() if duplicate['created_at'] is not None else '-')}"
            )
        return _result(message, answer, language, "human_escalation")

    row = db.execute(
        text(
            """
            INSERT INTO internship_escalations
                (internship_id, student_id, lecturer_id, escalation_type, severity,
                 target, subject, description, status)
            VALUES
                (:internship_id, :student_id, :lecturer_id, :escalation_type, :severity,
                 :target, :subject, :description, 'OPEN')
            RETURNING id
            """
        ),
        {
            "internship_id": int(internship["id"]),
            "student_id": student_id,
            "lecturer_id": lecturer_id,
            "escalation_type": draft.escalation_type,
            "severity": draft.severity,
            "target": draft.target,
            "subject": str(draft.subject),
            "description": str(draft.description),
        },
    ).mappings().first()
    escalation_id = int(row["id"])

    if draft.target == "FACULTY_MENTOR" and lecturer_id:
        db.execute(
            text(
                """
                INSERT INTO lecturer_student_messages
                    (lecturer_id, student_id, internship_id, message_type, content)
                VALUES
                    (:lecturer_id, :student_id, :internship_id, 'WARNING', :content)
                """
            ),
            {
                "lecturer_id": int(lecturer_id),
                "student_id": student_id,
                "internship_id": int(internship["id"]),
                "content": f"[Escalation #{escalation_id}] {draft.subject}: {draft.description}",
            },
        )
        db.execute(
            text(
                """
                INSERT INTO notifications
                    (user_id, title, message, notification_type, severity, related_type, related_id)
                VALUES
                    (:user_id, :title, :message, 'INTERNSHIP_ESCALATION', :severity, 'internship_escalation', :related_id)
                """
            ),
            {
                "user_id": int(lecturer_id),
                "title": draft.subject,
                "message": draft.description,
                "severity": "ERROR" if draft.severity in {"HIGH", "CRITICAL"} else "WARNING",
                "related_id": escalation_id,
            },
        )
    db.commit()

    target = (
        "Faculty Mentor"
        if draft.target == "FACULTY_MENTOR"
        else "CAID queue"
        if draft.target == "CAID_QUEUE"
        else str(draft.target)
    )

    if draft.target == "FACULTY_MENTOR":
        if lecturer_id:
            delivery_status_vi = (
                "Đã tạo bản ghi escalation, gửi một cảnh báo vào luồng Lecturer–Student "
                "và tạo notification cho Faculty Mentor được gán cho kỳ thực tập."
            )
            delivery_status_en = (
                "The escalation record was created, a warning was added to the Lecturer–Student "
                "message stream, and a notification was created for the assigned Faculty Mentor."
            )
        else:
            delivery_status_vi = (
                "Đã tạo bản ghi escalation nhưng kỳ thực tập hiện chưa có Faculty Mentor được gán, "
                "nên hệ thống chưa thể tạo notification cho người nhận."
            )
            delivery_status_en = (
                "The escalation record was created, but no Faculty Mentor is currently assigned "
                "to the internship, so no recipient notification could be created."
            )
    else:
        delivery_status_vi = (
            "Đã tạo bản ghi escalation với đích CAID queue. Hệ thống chưa gửi email tự động; "
            "record này sẽ được đọc qua luồng/dashboard xử lý escalation."
        )
        delivery_status_en = (
            "The escalation record was created for the CAID queue. No automatic email was sent; "
            "the record is available through the escalation handling flow/dashboard."
        )

    if language == "vi":
        answer = (
            f"Đã tạo escalation **#{escalation_id}**. Đây là chính xác những gì hệ thống đã ghi nhận:\n\n"
            f"- **Chủ đề:** {draft.subject or '-'}\n"
            f"- **Nội dung:** {draft.description}\n"
            f"- **Loại escalation:** {draft.escalation_type}\n"
            f"- **Mức độ:** {draft.severity}\n"
            f"- **Nơi nhận:** {target}\n"
            f"- **Trạng thái ban đầu:** OPEN\n"
            f"- **Internship ID:** {int(internship['id'])}\n\n"
            f"**Hệ thống đã thực hiện:** {delivery_status_vi}\n\n"
            "Bạn có thể hỏi **“cho tôi xem escalation vừa tạo”** để xem lại record này."
        )
    else:
        answer = (
            f"Escalation **#{escalation_id}** was created. This is exactly what the system recorded:\n\n"
            f"- **Subject:** {draft.subject or '-'}\n"
            f"- **Content:** {draft.description}\n"
            f"- **Escalation type:** {draft.escalation_type}\n"
            f"- **Severity:** {draft.severity}\n"
            f"- **Target:** {target}\n"
            f"- **Initial status:** OPEN\n"
            f"- **Internship ID:** {int(internship['id'])}\n\n"
            f"**What the system actually did:** {delivery_status_en}\n\n"
            "You can ask **“show me the escalation I just created”** to view it again."
        )

    return _result(message, answer, language, "human_escalation")